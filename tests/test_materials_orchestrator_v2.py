"""Regression tests for the bounded, resumable CV/CL audit contract."""

from __future__ import annotations

import json

import pytest

from docx import Document

from tools.io_utils import atomic_write_json, atomic_write_text
from tools.workflow.engine import dispatch
from tools.workflow.materials_contract import build_claim_contract
from tools.workflow.materials_hashes import metadata_hash, semantic_hash
from tools.workflow.materials_orchestrator import (
    MAX_AUDIT_ATTEMPTS,
    build_audit_task_packet,
    ensure_run,
    load_run,
    record_audit_result,
    reset,
    validate_audit_result,
)
from tools.workflow.materials_rules import build_rule_pack, rules_digest, validate_rule_pack
from tools.workflow.package_validator import _audit_receipt_matches
from tools.workflow.testing_packages import baseline_transform_fixture, build_package, build_workspace, canonical_fixture


pytestmark = pytest.mark.legacy


def _claim_contract(job_id: str = "C0-001"):
    return build_claim_contract(
        job_id=job_id,
        claim_ledger=[
            {
                "claim_id": "C1",
                "evidence_id": "EVID-AAA",
                "text": "Supported contract review for a payments team.",
                "kind": "Transferable",
                "assessment": "transferable",
            }
        ],
        forbidden_claims=["admitted as a solicitor"],
        evidence_allocation={"cv": ["C1"], "cover_letter": ["C1"]},
        jd_sha256="jd-1",
    )


def _write_cv_cl(package):
    cv = Document()
    cv.add_paragraph("Paralegal at Acme. Supported contract review for a payments team.")
    cv.save(package / "Candidate_CV.docx")
    cl = Document()
    cl.add_paragraph("Paralegal at Acme. I can bring bounded contract-review support.")
    cl.save(package / "Candidate_Cover Letter.docx")
    (package / "application_email.txt").write_text("Email remains outside the CV/CL child audit.", encoding="utf-8")
    atomic_write_json(package / "materials_draft.canonical.json", canonical_fixture("C0-001"))


def test_compiled_rules_are_small_and_p2_is_advisory():
    pack = build_rule_pack()
    assert pack["scope"] == "jd_mapping_and_presentation"
    assert pack["gate_policy"]["blocking_severities"] == ["P0", "P1"]
    assert pack["gate_policy"]["max_audit_attempts"] == 3
    assert pack["rules_digest"] == rules_digest(pack["rules"])
    assert any(item["severity"] == "P2" for item in pack["rules"])
    assert not validate_rule_pack(build_rule_pack(include_p2=False))


def test_semantic_hash_is_unchanged_by_docx_metadata_only_edit(tmp_path):
    path = tmp_path / "Candidate_CV.docx"
    document = Document()
    document.add_paragraph("Same body")
    document.core_properties.title = "Template Base v4"
    document.save(path)
    body_before = semantic_hash(path)
    metadata_before = metadata_hash(path)
    document = Document(path)
    document.core_properties.title = "Candidate CV"
    document.save(path)
    assert semantic_hash(path) == body_before
    assert metadata_hash(path) != metadata_before


def test_audit_packet_contains_jd_mapping_inputs_but_not_email_or_contract(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    contract = _claim_contract()
    run = ensure_run(package, job_id="C0-001", jd_text="Draft and review contracts.", claim_contract=contract)
    task = build_audit_task_packet(
        package,
        job_id="C0-001",
        jd_text="Draft and review contracts.",
        claim_contract=contract,
        producer_context_id="main-1",
    )
    assert task["audit_input_fingerprint"] == run["audit_input_fingerprint"]
    assert "claim_contract_sha256" not in run
    assert "claim_contract" not in task
    assert "application_email.txt" not in json.dumps(task, ensure_ascii=False)
    assert task["audit_scope"] == "jd_mapping_and_presentation"
    assert "claim_contract.json" not in json.dumps(task, ensure_ascii=False)
    assert task["forbidden"]


def test_repeat_finding_and_max_attempts_close_the_loop(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    contract = _claim_contract()
    run = ensure_run(package, job_id="C0-001", jd_text="Draft and review contracts.", claim_contract=contract)
    expected = {
        "job_id": "C0-001",
        "audit_input_fingerprint": run["audit_input_fingerprint"],
        "producer_context_id": "main-1",
    }
    report = {
        "job_id": "C0-001",
        "audit_scope": "jd_mapping_and_presentation",
        "audit_input_fingerprint": run["audit_input_fingerprint"],
        "auditor_context_id": "audit-1",
        "counts": {"P0": 0, "P1": 1, "P2": 0},
        "findings": [
            {
                "finding_id": "f1",
                "severity": "P1",
                "rule_id": "MAP-001",
                "material": "cover_letter",
                "quote": "generic wording",
                "reason": "does not answer the JD",
                "required_action": "add JD evidence",
            }
        ],
    }
    first = record_audit_result(package, report, expected=expected)
    assert first["status"] == "repair_required"
    assert load_run(package)["audit_attempts"] == 1
    second = record_audit_result(package, {**report, "auditor_context_id": "audit-2"}, expected=expected)
    assert second["status"] == "audit_loop_detected"
    assert load_run(package)["audit_attempts"] == 2

    # A different finding cannot escape the repeat-finding circuit breaker.
    other = {**report, "auditor_context_id": "audit-3", "findings": [{**report["findings"][0], "finding_id": "f2", "quote": "another generic sentence"}]}
    import pytest

    with pytest.raises(ValueError, match="audit_budget_exhausted"):
        record_audit_result(package, other, expected=expected)
    assert MAX_AUDIT_ATTEMPTS == 3


def test_p2_only_audit_passes_content_gate(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    contract = _claim_contract()
    run = ensure_run(package, job_id="C0-001", jd_text="Draft and review contracts.", claim_contract=contract)
    report = record_audit_result(
        package,
        {
            "job_id": "C0-001",
            "audit_scope": "jd_mapping_and_presentation",
            "audit_input_fingerprint": run["audit_input_fingerprint"],
            "auditor_context_id": "audit-1",
            "counts": {"P0": 0, "P1": 0, "P2": 1},
            "findings": [
                {
                    "finding_id": "p2",
                    "severity": "P2",
                    "rule_id": "OPT-001",
                    "material": "cover_letter",
                    "quote": "could be more differentiated",
                    "reason": "advisory",
                    "required_action": "optional",
                }
            ],
        },
        expected={
            "job_id": "C0-001",
            "audit_input_fingerprint": run["audit_input_fingerprint"],
            "producer_context_id": "main-1",
        },
    )
    assert report["status"] == "passed"
    assert report["content_ready_for_render"] is True
    assert report["ready_for_pdf"] is False


def test_audit_evidence_detects_post_audit_report_edit(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    contract = _claim_contract()
    run = ensure_run(package, job_id="C0-001", jd_text="Draft and review contracts.", claim_contract=contract)
    report = record_audit_result(
        package,
        {
            "job_id": "C0-001",
            "audit_scope": "jd_mapping_and_presentation",
            "audit_input_fingerprint": run["audit_input_fingerprint"],
            "auditor_context_id": "audit-1",
            "counts": {"P0": 0, "P1": 0, "P2": 0},
            "findings": [],
        },
        expected={"job_id": "C0-001", "audit_input_fingerprint": run["audit_input_fingerprint"], "producer_context_id": "main-1"},
    )
    assert _audit_receipt_matches(package, report, {}) is True
    report["findings"] = [{"severity": "P2", "rule_id": "OPT-001", "material": "cv"}]
    (package / "materials_audit.json").write_text(json.dumps(report), encoding="utf-8")
    tampered = json.loads((package / "materials_audit.json").read_text(encoding="utf-8"))
    assert _audit_receipt_matches(package, tampered, {}) is False


def test_child_cannot_self_resolve_a_finding(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    contract = _claim_contract()
    run = ensure_run(package, job_id="C0-001", jd_text="Draft and review contracts.", claim_contract=contract)
    import pytest

    with pytest.raises(ValueError, match="finding_status_invalid"):
        record_audit_result(
            package,
            {
                "job_id": "C0-001",
                "audit_scope": "jd_mapping_and_presentation",
                "audit_input_fingerprint": run["audit_input_fingerprint"],
                "auditor_context_id": "audit-1",
                "counts": {"P0": 0, "P1": 0, "P2": 0},
                "findings": [{
                    "finding_id": "f1", "severity": "P1", "rule_id": "MAP-001",
                    "material": "cover_letter", "status": "repaired", "quote": "x",
                    "reason": "y", "required_action": "z",
                }],
            },
            expected={"job_id": "C0-001", "audit_input_fingerprint": run["audit_input_fingerprint"], "producer_context_id": "main-1"},
        )


def test_child_format_finding_is_rejected_even_when_tagged_as_cv(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    run = ensure_run(package, job_id="C0-001", jd_text="Draft and review contracts.")
    task = build_audit_task_packet(
        package,
        job_id="C0-001",
        jd_text="Draft and review contracts.",
        producer_context_id="main-1",
    )
    report = {
        "job_id": "C0-001",
        "audit_scope": "jd_mapping_and_presentation",
        "audit_input_fingerprint": run["audit_input_fingerprint"],
        "auditor_context_id": task["auditor_context_id"],
        "counts": {"P0": 0, "P1": 1, "P2": 0},
        "findings": [{
            "finding_id": "format-1",
            "severity": "P1",
            "rule_id": "MAP-001",
            "material": "cv",
            "target_id": "cv-heading",
            "quote": "The PDF has two pages",
            "reason": "page count is not a content finding",
            "required_action": "leave this to the host format gate",
        }],
    }
    errors = validate_audit_result(
        report,
        expected={
            "job_id": "C0-001",
            "audit_input_fingerprint": run["audit_input_fingerprint"],
            "producer_context_id": "main-1",
            "auditor_context_id": task["auditor_context_id"],
            "task": task,
        },
    )
    assert "audit_scope_contains_format_finding" in errors
    report["ready_for_pdf"] = True
    envelope_errors = validate_audit_result(report, expected={
        "job_id": "C0-001",
        "audit_input_fingerprint": run["audit_input_fingerprint"],
        "producer_context_id": "main-1",
        "auditor_context_id": task["auditor_context_id"],
        "task": task,
    })
    assert "audit_scope_contains_format_output" in envelope_errors


def test_child_task_cannot_fall_back_to_legacy_docx_or_pdf(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    cv = Document()
    cv.add_paragraph("Legacy CV text")
    cv.save(package / "Legacy_CV.docx")
    import pytest

    with pytest.raises(ValueError, match="canonical_draft_required_before_content_audit"):
        build_audit_task_packet(
            package,
            job_id="C0-001",
            jd_text="Draft and review contracts.",
            producer_context_id="main-1",
        )


def test_auditor_cannot_request_negative_self_disclosure_as_map_repair(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    _write_cv_cl(package)
    task = build_audit_task_packet(
        package,
        job_id="C0-001",
        jd_text="Cantonese is required for client calls.",
        producer_context_id="main-1",
    )
    report = {
        "job_id": "C0-001",
        "audit_scope": "jd_mapping_and_presentation",
        "audit_input_fingerprint": task["audit_input_fingerprint"],
        "auditor_context_id": task["auditor_context_id"],
        "counts": {"P0": 0, "P1": 1, "P2": 0},
        "findings": [
            {
                "finding_id": "gap-1",
                "severity": "P1",
                "rule_id": "MAP-001",
                "material": "cover_letter",
                "target_id": "cl-opening",
                "quote": "Cantonese is required",
                "reason": "The requirement has no positive evidence.",
                "required_action": "State that Cantonese is not declared in my language profile.",
            }
        ],
    }

    errors = validate_audit_result(
        report,
        expected={
            "job_id": "C0-001",
            "audit_input_fingerprint": task["audit_input_fingerprint"],
            "producer_context_id": "main-1",
            "auditor_context_id": task["auditor_context_id"],
            "task": task,
        },
    )

    assert "audit_repair_requests_negative_disclosure" in errors


def test_reset_is_preview_then_archives_old_audit(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    atomic_write_text(package / "materials_audit.md", "old audit\n")
    preview = reset(package, scope="audit")
    assert preview["status"] == "preview"
    assert (package / "materials_audit.md").is_file()
    done = reset(package, scope="audit", confirm=True)
    assert done["status"] == "reset"
    assert not (package / "materials_audit.md").exists()
    assert list((package / ".history").rglob("materials_audit.md"))


def test_strict_audit_gateway_creates_task_and_fails_closed(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    _write_cv_cl(package)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "model_plan": plan})["status"] == "succeeded"
    drafted = dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "canonical_draft": baseline_transform_fixture(package)})
    assert drafted["status"] == "succeeded"
    out = dispatch("audit", workspace=ws, payload={"job_id": "C0-001", "strict": True})
    assert out["status"] == "blocked"
    assert "independent_audit_required" in out["blockers"]
    assert (package / "materials_audit_task.json").is_file()
    auto = dispatch("audit", workspace=ws, payload={"job_id": "C0-001", "strict": True, "auto_audit": True})
    assert auto["status"] == "blocked"
    assert auto["audit_dispatch"]["status"] == "delegation_required"
    assert auto["audit_dispatch"]["confirmation_required"] is False


def test_docx_is_created_only_after_content_audit_and_metadata_is_sanitized(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "model_plan": plan})["status"] == "succeeded"
    drafted = dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "canonical_draft": baseline_transform_fixture(package)})
    assert not list(package.glob("*.docx"))
    task = drafted["audit_task_packet"]
    report = {"job_id": "C0-001", "audit_scope": "jd_mapping_and_presentation", "audit_input_fingerprint": task["audit_input_fingerprint"], "auditor_context_id": task["auditor_context_id"], "counts": {"P0": 0, "P1": 0, "P2": 0}, "findings": []}
    assert dispatch("audit", workspace=ws, payload={"job_id": "C0-001", "audit_result": report})["status"] == "succeeded"
    assert dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "stage": "render"})["status"] == "succeeded"
    for path in package.glob("*.docx"):
        assert "base v4" not in (Document(path).core_properties.title or "").casefold()


def test_content_audit_repair_can_reenter_drafting_without_skipping_gate(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "model_plan": plan})["status"] == "succeeded"
    drafted = dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "canonical_draft": baseline_transform_fixture(package)})
    assert drafted["status"] == "succeeded"
    run = load_run(package)
    task = json.loads((package / "materials_audit_task.json").read_text(encoding="utf-8"))
    report = {
        "job_id": "C0-001",
        "audit_scope": "jd_mapping_and_presentation",
        "audit_input_fingerprint": run["audit_input_fingerprint"],
        "auditor_context_id": task["auditor_context_id"],
        "producer_context_id": task["producer_context_id"],
        "counts": {"P0": 0, "P1": 1, "P2": 0},
        "findings": [{
            "finding_id": "f1", "severity": "P1", "rule_id": "MAP-001",
            "material": "cover_letter", "quote": "generic", "reason": "no JD answer",
            "required_action": "add evidence",
        }],
    }
    changed_cl = next(
        item
        for item in task["tailoring_delta"]["changes"]
        if item["material"] == "cover_letter"
    )
    report["findings"][0]["quote"] = changed_cl["after"]
    report["findings"][0]["target_id"] = changed_cl["target_id"]
    blocked = dispatch("audit", workspace=ws, payload={"job_id": "C0-001", "audit_result": report})
    assert blocked["status"] == "blocked"
    current = json.loads((package / "materials_draft.canonical.json").read_text(encoding="utf-8"))
    before = next(
        item["text"]
        for item in current["cover_letter"]["blocks"]
        if item["id"] == changed_cl["target_id"]
    )
    retry = dispatch("materials", workspace=ws, payload={"job_id": "C0-001", "repair_patch": {
        "job_id": "C0-001", "base_canonical_sha256": current["canonical_sha256"],
        "audit_input_fingerprint": run["audit_input_fingerprint"],
        "changes": [{"finding_ids": ["f1"], "material": "cover_letter", "target_id": changed_cl["target_id"], "before_text": before, "after_text": before + " I can provide accurate, bounded support."}],
    }})
    assert retry["status"] == "succeeded"
    assert retry["after_state"] == "content_audit_pending"
    assert retry["audit_dispatch"]["confirmation_required"] is False
