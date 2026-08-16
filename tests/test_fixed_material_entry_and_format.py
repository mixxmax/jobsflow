"""Hard boundaries for the product materials entry and format chain."""

from __future__ import annotations

import csv
import json

import pytest

from tools.job_materials.packages import create_package_from_tracker
from tools.workflow.engine import dispatch
from tools.workflow.materials_renderer import render_canonical_docx
from tools.workflow.materials_validator import validate_materials_packet
from tools.workflow.package_context import PackageContextLoader
from tools.workflow.testing_packages import build_package, build_workspace, prepare_package_for_apply


def _tracker(ws, row: dict[str, str]) -> None:
    path = ws / "02_Tracker" / "hk_apply_list_2026-08-15.csv"
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = ["岗位编号", "职位", "公司", "层级", "简历版本", "链接", "来源"]
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerow({field: row.get(field, "") for field in fields})


def test_renderer_is_bound_to_lane_master_styles(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)

    from docx import Document

    from tools.workflow.materials_renderer import expected_filenames

    names = expected_filenames(package, ws)
    cv = Document(package / names["cv_docx"])
    cl = Document(package / names["cl_docx"])
    assert any(p.style.name == "Resume Bullet" for p in cv.paragraphs)
    assert any(p.style.name == "Letter Body" for p in cl.paragraphs)
    assert {style.name for style in cv.styles} >= {"Resume Section", "Job Heading", "Resume Bullet", "Compact Line"}
    receipt = json.loads((package / "materials_render_receipt.json").read_text(encoding="utf-8"))
    assert receipt["renderer_version"] != "canonical-docx-v1"
    assert receipt["template_paths"]["cv"]
    assert receipt["template_sha256"]["cv"]


def test_renderer_bounds_long_company_and_role_filename(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    manifest = json.loads((package / "job_manifest.json").read_text(encoding="utf-8"))
    legal_company = "Industrial and Commercial Bank of China (Asia) Limited"
    manifest["job"].update(
        {
            "publisher_type": "employer",
            "publisher_name": legal_company,
            "company_out": legal_company,
            "employer_name": legal_company,
            "role_material": "Officer to Assistant Manager, CDD, Channel Management Dept",
            "role_display": "Officer to Assistant Manager, CDD, Channel Management Dept",
        }
    )
    (package / "job_manifest.json").write_text(json.dumps(manifest), encoding="utf-8")

    from tools.workflow.materials_renderer import expected_filenames

    names = expected_filenames(package, ws)
    assert names["cv_pdf"] == "Test Candidate ICBC Asia CDD Channel Management CV.pdf"
    assert len(names["cv_pdf"].removesuffix(".pdf").removesuffix(" CV")) <= 80


def test_renderer_refuses_to_fall_back_to_plain_document(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)
    for path in (ws / "01_Masters" / "C_track").glob("*.docx"):
        path.unlink()
    with pytest.raises(ValueError, match="base_template_missing"):
        render_canonical_docx(package, ws, force=True)


def test_direct_package_docx_to_pdf_conversion_is_not_a_second_entry(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    from docx import Document

    plain = Document()
    plain.add_paragraph("plain text")
    path = package / "Test Candidate_Acme_Paralegal CV.docx"
    plain.save(path)
    from tools.fresh_24h.docx_to_pdf import convert

    with pytest.raises(RuntimeError, match="fixed_material_entry_required"):
        convert(path, package / "plain.pdf")


def test_legacy_auto_compress_cannot_mutate_package_docx(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)
    from docx import Document

    path = package / "Test Candidate_Acme_Paralegal CV.docx"
    document = Document()
    document.add_paragraph("plain text")
    document.save(path)
    from tools.job_materials.auto_compress import auto_compress

    with pytest.raises(RuntimeError, match="fixed_material_entry_required"):
        auto_compress(path, dry_run=True)


def test_materials_cannot_create_a_package_after_entry_is_missing(tmp_path):
    ws = build_workspace(tmp_path)
    _tracker(
        ws,
        {
            "岗位编号": "C0-071",
            "职位": "Compliance Analyst",
            "公司": "Acme",
            "层级": "核心",
            "简历版本": "C",
            "链接": "https://example.test/71",
            "来源": "LinkedIn",
        },
    )
    context = PackageContextLoader(ws).load("C0-071")
    assert context.package is None
    assert "package_missing" in context.blockers
    assert not list((ws / "01_Masters").rglob("C0-071_*"))


def test_package_path_is_derived_from_id_and_lane_mismatch_is_blocked(tmp_path):
    ws = build_workspace(tmp_path)
    _tracker(
        ws,
        {
            "岗位编号": "C1-065",
            "职位": "KYC Officer",
            "公司": "Peoplebank",
            "层级": "一级",
            "简历版本": "F",
            "链接": "https://example.test/65",
            "来源": "JobsDB",
        },
    )
    with pytest.raises(ValueError, match="lane_binding_mismatch"):
        create_package_from_tracker(ws, "C1-065")
    assert not list((ws / "01_Masters" / "F_track").rglob("C1-065_*"))


def test_entry_package_contains_an_immutable_path_binding(tmp_path):
    ws = build_workspace(tmp_path)
    _tracker(
        ws,
        {
            "岗位编号": "C1-072",
            "职位": "Compliance Analyst",
            "公司": "Acme",
            "层级": "一级",
            "简历版本": "C",
            "链接": "https://example.test/72",
            "来源": "LinkedIn",
        },
    )
    package = create_package_from_tracker(ws, "C1-072")
    assert package == (ws / "01_Masters" / "C_track" / "一级" / "C1-072_未投_Acme").resolve()
    binding = json.loads((package / "package_binding.json").read_text(encoding="utf-8"))
    assert binding["job_id"] == "C1-072"
    assert binding["lane"] == "C"
    assert binding["tier"]["code"] == "1"
    assert binding["expected_relative_path"] == "01_Masters/C_track/一级/C1-072_未投_Acme"


def test_confirmed_push_creates_the_bound_package_before_materials(tmp_path, monkeypatch):
    ws = build_workspace(tmp_path)
    scored = ws / "02_Tracker" / "scored.csv"
    scored.write_text(
        "职位,公司,链接,CareerOps分数,CareerOps等级,简历版本\n"
        "Compliance Analyst,Acme,https://example.test/73,4.2,A,C\n",
        encoding="utf-8",
    )
    from tools.workflow.adapters.scan import write_run_record

    write_run_record(ws, run_id="run-entry-package", mode="temp", scored_path=scored)
    monkeypatch.setenv("JOBSFlow_FRESH_BACKEND", "csv")
    preview = dispatch("push", workspace=ws, payload={"run_id": "run-entry-package"})
    assert preview["status"] == "planned"
    pushed = dispatch(
        "push",
        workspace=ws,
        payload={
            "run_id": "run-entry-package",
            "confirmation_id": preview["proposal_id"],
        },
    )
    assert pushed["status"] == "succeeded"
    assert pushed["package_paths"]
    package = next((ws / "01_Masters").rglob("*-001_未投_Acme"))
    assert package.is_dir()
    assert (package / "package_binding.json").is_file()
    assert PackageContextLoader(ws).load(package.name.split("_", 1)[0]).package == str(package)


def test_negative_qualification_self_disclosure_is_blocked(tmp_path):
    packet = {
        "full_jd": "The role values clear client communication.",
        "facts": [{"id": "EVID-1", "text": "Managed client records."}],
        "assessment": {"match_type": "transferable"},
        "preflight": {},
        "evidence_ids": ["EVID-1"],
        "outbound": {
            "cv_text": "Managed client records.",
            "cl_text": "Cantonese is not declared in my language profile.",
        },
    }
    result = validate_materials_packet(packet)
    assert "negative_self_disclosure" in {item["code"] for item in result["errors"]}
