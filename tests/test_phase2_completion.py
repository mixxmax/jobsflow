"""Completion-seam tests for the second workflow phase.

These tests exercise public workflow seams rather than private helper calls.
They intentionally cover the failure modes that a green module-only suite can
miss: the scan watermark, the real local push sink, and the material audit
receipt.
"""

from __future__ import annotations

import json
from pathlib import Path

from tools.io_utils import atomic_write_json
from tools.workflow.adapters import push as push_adapter
from tools.workflow.engine import dispatch
from tools.workflow.entity_state import load_entity_state
from tools.workflow.fresh_store import LocalCsvFreshStore, default_fresh_store
from tools.workflow.package_validator import MaterialsPackageValidator
from tools.workflow.testing_packages import build_package, build_workspace, prepare_package_for_apply, canonical_fixture, write_minimal_pdf


def _scored_run(ws: Path, run_id: str = "run-completion") -> None:
    tracker = ws / "02_Tracker"
    tracker.mkdir(parents=True, exist_ok=True)
    scored = tracker / "fresh_24h_2026-08-14_twopass_scored.csv"
    scored.write_text(
        "岗位编号,职位,公司,链接,CareerOps分数,评估状态\n"
        "C0-901,Role,Acme,https://example.test/901,4.2,\n",
        encoding="utf-8",
    )
    from tools.workflow.adapters.scan import write_run_record

    write_run_record(ws, run_id=run_id, mode="temp", scored_path=scored)


def test_workflow_push_uses_a_real_local_csv_sink(tmp_path, monkeypatch):
    ws = build_workspace(tmp_path)
    _scored_run(ws)
    monkeypatch.setenv("JOBSFlow_FRESH_BACKEND", "csv")
    preview = dispatch("push", workspace=ws, payload={"run_id": "run-completion"})
    assert preview["status"] == "planned"
    out = dispatch(
        "push",
        workspace=ws,
        payload={"run_id": "run-completion", "confirmation_id": preview["proposal_id"]},
    )
    assert out["status"] == "succeeded"
    assert out["backend"] == "local_csv"
    assert out["postconditions"] == ["fresh_rows_read_back"]
    assert list((ws / "02_Tracker").glob("fresh_24h_*.csv"))
    assert not (ws / "02_Tracker" / "workflow" / "fresh" / "fresh_24h_run-completion" / "active.json").exists()


def test_auto_backend_fails_safe_to_durable_csv_without_sheet_credentials(tmp_path, monkeypatch):
    monkeypatch.delenv("GSHEET_ID", raising=False)
    monkeypatch.delenv("GOOGLE_APPLICATION_CREDENTIALS", raising=False)
    monkeypatch.delenv("JOBSFlow_FRESH_BACKEND", raising=False)
    store = default_fresh_store(tmp_path / "JobSearch_2026", "fresh_24h_2026-08-14")
    assert isinstance(store, LocalCsvFreshStore)
    store.merge_incoming([{"岗位编号": "C0-902", "职位": "Role", "公司": "Acme", "链接": "https://example.test/902"}])
    assert store.read_active().row_count == 1
    assert store.active_path.is_file()


def test_refresh_commit_uses_scan_window_until_not_scoring_finish(tmp_path):
    ws = build_workspace(tmp_path)
    tracker = ws / "02_Tracker"
    state_path = tracker / "fresh_refresh_state.json"
    atomic_write_json(
        state_path,
        {
            "version": 1,
            "last_refresh_at": "2026-08-14T09:00:00Z",
            "last_mode": "temp",
            "last_window_hours": 2,
            "history": [],
        },
    )
    scored = tracker / "fresh_24h_2026-08-14_twopass_scored.csv"
    scored.write_text("岗位编号,链接\nC0-901,https://example.test/901\n", encoding="utf-8")
    summary = tracker / "fresh_24h_2026-08-14_run.json"
    atomic_write_json(
        summary,
        {
            "mode": "temp",
            "hours": 2,
            "day": "2026-08-14",
            "window": {
                "since": "2026-08-14T09:00:00Z",
                "until": "2026-08-14T11:00:00Z",
            },
            "counts": {"new": 1},
            "candidates_csv": str(tracker / "fresh_24h_2026-08-14.csv"),
        },
    )
    from tools.workflow.refresh_commit import commit_refresh_after_score

    result = commit_refresh_after_score(workspace=ws, mode="temp")
    assert result is not None
    state = json.loads(state_path.read_text(encoding="utf-8"))
    assert state["last_refresh_at"] == "2026-08-14T11:00:00Z"


def test_generic_outbound_without_role_or_employer_is_not_apply_ready(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    for path in package.glob("*.txt"):
        path.unlink()
    (package / "Pat_Paralegal_Acme_CV.docx").unlink(missing_ok=True)
    (package / "Pat_Paralegal_Acme_Cover_Letter.docx").unlink(missing_ok=True)

    from docx import Document
    from tools.io_utils import atomic_write_text
    from tools.workflow.testing_packages import write_minimal_pdf

    cv = Document()
    cv.add_paragraph("Reviewed contracts and maintained reliable records.")
    cv.save(package / "Pat_CV_Acme.docx")
    cl = Document()
    cl.add_paragraph("I can bring adjacent experience to this opportunity.")
    cl.save(package / "Pat_Cover_Letter_Acme.docx")
    write_minimal_pdf(package / "Pat_CV_Acme.pdf", "Reviewed contracts")
    write_minimal_pdf(package / "Pat_Cover_Letter_Acme.pdf", "Adjacent experience")
    atomic_write_text(package / "application_email.md", "Please find my materials.\n")

    report = MaterialsPackageValidator().validate(package)
    assert report["apply_ready"] is False
    assert "role_or_employer_missing" in {item["code"] for item in report["findings"]}


def test_apply_requires_a_hash_bound_audit_receipt(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)
    (package / "materials_audit.json").unlink()
    report = MaterialsPackageValidator().validate(package)
    assert report["apply_ready"] is False
    assert "content_audit_missing" in {item["code"] for item in report["findings"]}


def test_forged_receipt_without_independent_audit_is_rejected(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)
    (package / "materials_audit.json").unlink()
    atomic_write_json(
        package / "materials_audit.json",
        {
            "status": "passed",
            "auditor": "model",
            "artifact_hashes": {},
        },
    )
    report = MaterialsPackageValidator().validate(package)
    assert report["apply_ready"] is False
    assert "content_audit_missing" in {item["code"] for item in report["findings"]}


def test_public_material_lifecycle_reaches_apply_ready_state(tmp_path):
    ws = build_workspace(tmp_path)
    build_package(ws)
    prepare_package_for_apply(ws)
    out = dispatch("apply", workspace=ws, payload={"job_id": "C0-001"})
    assert out["status"] == "succeeded"
    assert out["apply_ready"] is True
    assert out["submitted"] is False
    state = load_entity_state(ws, "materials", "C0-001")
    assert state.phase == "apply_ready"


def test_cli_entrypoints_drive_the_complete_synthetic_chain(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    fixture = tmp_path / "scan_fixture.json"
    atomic_write_json(
        fixture,
        {
            "run_id": "run-cli",
            "jobs": [{"job_id": "C0-001", "title": "Paralegal", "company": "Acme", "score": "4.2"}],
        },
    )
    plan = package / "materials_plan.validated.json"
    from tools.workflow.__main__ import main

    assert main(["scan", "--workspace", str(ws), "--fixture", str(fixture)]) == 0
    preview = dispatch(
        "push",
        workspace=ws,
        payload={"run_id": "run-cli", "backend": "csv"},
    )
    assert preview["status"] == "planned"
    assert main(
        [
            "push",
            "--workspace",
            str(ws),
            "--run-id",
            "run-cli",
            "--backend",
            "csv",
            "--confirm",
            preview["proposal_id"],
        ]
    ) == 0
    assert main(["materials", "--workspace", str(ws), "--job-id", "C0-001", "--plan", str(plan)]) == 0
    canonical = tmp_path / "canonical.json"
    atomic_write_json(canonical, canonical_fixture())
    assert main(["materials", "draft", "--workspace", str(ws), "--job-id", "C0-001", "--content", str(canonical)]) == 0
    task = json.loads((package / "materials_audit_task.json").read_text(encoding="utf-8"))
    audit_result = tmp_path / "audit_result.json"
    atomic_write_json(audit_result, {"job_id": "C0-001", "audit_scope": "jd_mapping_and_presentation", "audit_input_fingerprint": task["audit_input_fingerprint"], "auditor_context_id": task["auditor_context_id"], "counts": {"P0": 0, "P1": 0, "P2": 0}, "findings": []})
    assert main(["audit", "--workspace", str(ws), "--job-id", "C0-001", "--result", str(audit_result)]) == 0
    assert main(["materials", "render", "--workspace", str(ws), "--job-id", "C0-001"]) == 0
    from tools.workflow.materials_renderer import expected_filenames
    names = expected_filenames(package, ws)
    write_minimal_pdf(package / names["cv_pdf"], "Paralegal contract review support and accurate records with IELTS 7.5 English experience.")
    write_minimal_pdf(package / names["cl_pdf"], "Application for Paralegal at Acme with contract review support and IELTS 7.5 English experience.")
    assert main(["materials", "--workspace", str(ws), "--job-id", "C0-001", "--stage", "pdf_generated"]) == 0
    assert main(["format", "--workspace", str(ws), "--job-id", "C0-001"]) == 0
    assert main(["apply", "--workspace", str(ws), "--job-id", "C0-001"]) == 0


def test_valid_package_cannot_bypass_material_state_chain(tmp_path):
    ws = build_workspace(tmp_path)
    build_package(ws)
    out = dispatch("apply", workspace=ws, payload={"job_id": "C0-001"})
    assert out["apply_ready"] is False
    assert out["status"] == "blocked"
    assert "content_audit_missing" in out["blockers"]
    assert load_entity_state(ws, "materials", "C0-001").phase == "idle"
