"""Product contracts for baseline-first CV/Cover Letter tailoring."""

from __future__ import annotations

import json

import pytest

from docx import Document

from tools.io_utils import atomic_write_json
from tools.workflow.engine import dispatch
from tools.workflow.materials_baseline import content_baseline_digest, load_content_baseline
from tools.workflow.materials_draft import load_canonical_draft
from tools.workflow.materials_drafting_context import load_drafting_scope
from tools.workflow.materials_hashes import audit_input_fingerprint
from tools.workflow.materials_rules import build_rule_pack
from tools.workflow.testing_packages import (
    build_package,
    build_workspace,
    canonical_fixture,
    prepare_package_for_apply,
)


def _drafting_binding(package):
    scope = load_drafting_scope(package, phase="tailoring")
    return {
        "drafting_context_id": scope["context_id"],
        "drafting_input_fingerprint": scope["input_fingerprint"],
    }


def test_materials_packet_uses_lane_masters_as_content_baseline(tmp_path):
    workspace = build_workspace(tmp_path)
    build_package(workspace, with_outbound=False)

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001"},
    )

    assert outcome["status"] == "succeeded"
    baseline = outcome["task_packet"]["content_baseline"]
    assert baseline["contract"]["unmentioned_blocks"] == "retain"
    assert baseline["contract"]["deletion_allowed"] is False
    assert any(
        block["text"] == "A measurable action and result."
        for block in baseline["cv"]["blocks"]
    )
    assert any(
        "Supported the relevant work." in block["text"]
        for block in baseline["cover_letter"]["blocks"]
    )
    assert baseline["baseline_sha256"]


def test_host_removes_optional_address_slot_and_seeds_jd_placeholders(tmp_path):
    workspace = build_workspace(tmp_path)
    cl_master = workspace / "01_Masters" / "C_track" / "cl_master_C_test_v1.docx"
    document = Document(str(cl_master))
    document.add_paragraph("[Company address]", style="Letter Compact")
    document.add_paragraph(
        "The role's focus on [principal priority from the job description] is relevant to my background.",
        style="Letter Body",
    )
    document.save(str(cl_master))
    package = build_package(workspace, with_outbound=False)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )

    assert outcome["status"] == "succeeded"
    baseline = outcome["task_packet"]["content_baseline"]
    address = next(block for block in baseline["cover_letter"]["blocks"] if block["text"] == "[Company address]")
    assert address["content_floor"] is False
    draft = load_canonical_draft(package)
    cl_text = "\n".join(block["text"] for block in draft["cover_letter"]["blocks"])
    assert "[Company address]" not in cl_text
    assert "[principal priority" not in cl_text
    assert "Draft vendor contracts" in cl_text


@pytest.mark.legacy
def test_host_owns_cover_letter_company_line_and_model_cannot_rewrite_it(tmp_path):
    workspace = build_workspace(tmp_path)
    cl_master = workspace / "01_Masters" / "C_track" / "cl_master_C_test_v1.docx"
    document = Document(str(cl_master))
    document.add_paragraph("Hiring Manager", style="Letter Compact")
    document.add_paragraph("[Company]", style="Letter Compact")
    document.save(str(cl_master))
    package = build_package(workspace, with_outbound=False, publisher_type="employer", publisher_name="Acme")
    manifest = json.loads((package / "job_manifest.json").read_text(encoding="utf-8"))
    manifest["job"]["employer_name"] = "Acme"
    manifest["job"]["company_out"] = "Acme"
    atomic_write_json(package / "job_manifest.json", manifest)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))

    planned = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )
    assert planned["status"] == "succeeded"
    assert planned["task_packet"]["cover_letter_header_contract"]["model_may_edit"] is False
    draft = load_canonical_draft(package)
    recipient = [
        block
        for block in draft["cover_letter"]["blocks"]
        if block.get("section") == "recipient"
    ]
    assert any(block["text"] == "Acme" for block in recipient)

    scope = _drafting_binding(package)
    company_block = next(block for block in (load_content_baseline(package)["cover_letter"]["blocks"]) if block.get("text") == "[Company]")
    blocked = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                "job_id": "C0-001",
                "baseline_sha256": load_content_baseline(package)["baseline_sha256"],
                **scope,
                "changes": [
                    {
                        "material": "cover_letter",
                        "baseline_id": company_block["id"],
                        "action": "rewrite",
                        "text": "Hiring Manager",
                        "jd_anchor_ids": ["JD-001"],
                    }
                ],
                "additions": [],
            },
        },
    )
    assert blocked["status"] == "blocked"
    assert "baseline_transform_host_managed_block" in blocked["error"]


def test_undisclosed_recruiter_gets_neutral_company_line_not_publisher(tmp_path):
    workspace = build_workspace(tmp_path)
    cl_master = workspace / "01_Masters" / "C_track" / "cl_master_C_test_v1.docx"
    document = Document(str(cl_master))
    document.add_paragraph("[Company]", style="Letter Compact")
    document.save(str(cl_master))
    package = build_package(workspace, with_outbound=False, publisher_type="recruiter", publisher_name="Adecco")
    manifest = json.loads((package / "job_manifest.json").read_text(encoding="utf-8"))
    manifest["job"]["company_out"] = ""
    manifest["job"]["employer_name"] = ""
    atomic_write_json(package / "job_manifest.json", manifest)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )
    assert outcome["status"] == "succeeded"
    cl_text = "\n".join(block["text"] for block in load_canonical_draft(package)["cover_letter"]["blocks"])
    assert "the hiring organisation" in cl_text
    assert "Adecco" not in cl_text


def test_unresolved_candidate_placeholder_is_not_replaced_with_jd_text(tmp_path):
    workspace = build_workspace(tmp_path)
    cv_master = workspace / "01_Masters" / "C_track" / "master_C_test_v1.docx"
    document = Document(str(cv_master))
    document.add_paragraph("[YOUR_EXPERIENCE]", style="Resume Bullet")
    document.save(str(cv_master))
    package = build_package(workspace, with_outbound=False)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )

    assert outcome["status"] == "blocked"
    assert "canonical_seed_invalid" in outcome["blockers"]
    assert "canonical_placeholder" in outcome["error"]


def test_drafting_contract_requests_only_a_bounded_delta(tmp_path):
    workspace = build_workspace(tmp_path)
    build_package(workspace, with_outbound=False)

    packet = dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001"}
    )["task_packet"]

    schema = packet["draft_seed_schema"]
    assert schema["artifact_type"] == "jobsflow_baseline_transform"
    assert schema["baseline_sha256"] == packet["content_baseline"]["baseline_sha256"]
    assert schema["unmentioned_blocks"] == "retain"
    assert schema["deletion_allowed"] is False
    assert "complete CV and Cover Letter" not in json.dumps(packet, ensure_ascii=False)


def test_plan_response_never_advertises_the_legacy_full_document_schema(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )

    assert outcome["status"] == "succeeded"
    assert outcome["draft_schema"]["artifact_type"] == "jobsflow_baseline_transform"
    assert outcome["draft_schema"]["unmentioned_blocks"] == "retain"
    assert outcome["draft_schema"]["deletion_allowed"] is False
    assert outcome["draft_schema"]["jd_anchor_catalog"][0] == {
        "id": "JD-001",
        "text": "Draft vendor contracts",
        "priority": 1,
    }
    assert "complete CV" not in json.dumps(outcome["draft_schema"], ensure_ascii=False)


def test_transform_rejects_an_anchor_not_present_in_the_validated_plan(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"
    cv_target = next(block for block in baseline["cv"]["blocks"] if block["type"] == "bullet")
    cl_target = next(block for block in baseline["cover_letter"]["blocks"] if block["type"] == "bullet")

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                "job_id": "C0-001",
                "baseline_sha256": baseline["baseline_sha256"],
                **_drafting_binding(package),
                "changes": [
                    {
                        "material": "cv",
                        "baseline_id": cv_target["id"],
                        "action": "rewrite",
                        "text": "JD-aligned contract review evidence.",
                        "jd_anchor_ids": ["JD-999"],
                    },
                    {
                        "material": "cover_letter",
                        "baseline_id": cl_target["id"],
                        "action": "rewrite",
                        "text": "JD-aligned contract review value.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                ],
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "blocked"
    assert "baseline_transform_jd_anchor_unknown" in outcome["error"]


def test_audit_fingerprint_binds_the_frozen_lane_content_baseline(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})
    pack = build_rule_pack()
    first = audit_input_fingerprint(
        package=package,
        jd_text="selected JD",
        rules_digest=pack["rules_digest"],
    )
    path = package / "materials_baseline.json"
    baseline = json.loads(path.read_text(encoding="utf-8"))
    baseline["cv"]["blocks"][-1]["text"] += " Stable evidence changed."
    baseline["baseline_sha256"] = content_baseline_digest(baseline)
    atomic_write_json(path, baseline)

    second = audit_input_fingerprint(
        package=package,
        jd_text="selected JD",
        rules_digest=pack["rules_digest"],
    )

    assert second != first


@pytest.mark.legacy
def test_apply_fails_closed_if_the_bound_lane_baseline_changes_after_render(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace)
    prepare_package_for_apply(workspace)
    path = package / "materials_baseline.json"
    baseline = json.loads(path.read_text(encoding="utf-8"))
    baseline["cv"]["blocks"][-1]["text"] += " Stable evidence changed."
    baseline["baseline_sha256"] = content_baseline_digest(baseline)
    atomic_write_json(path, baseline)

    outcome = dispatch("apply", workspace=workspace, payload={"job_id": "C0-001"})

    assert outcome["status"] == "blocked"
    assert "baseline_content_floor_invalid" in outcome["blockers"]


def test_audit_rules_review_tailoring_delta_without_freezing_baseline_wording():
    pack = build_rule_pack()
    rule = next(item for item in pack["rules"] if item["rule_id"] == "BASE-001")

    assert rule["severity"] == "P1"
    assert "rewrite" in rule["check"].casefold()
    assert "merge" in rule["check"].casefold()
    assert "jd" in rule["check"].casefold()
    assert "verbatim" not in rule["check"].casefold()


def test_untailored_baseline_cannot_skip_directly_to_content_audit(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"

    outcome = dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "stage": "drafting"}
    )

    assert outcome["status"] == "blocked"
    assert "baseline_transform_required" in outcome["blockers"]
    assert outcome["next_action"] == "submit_bounded_baseline_transform"


def test_transform_must_tailor_both_cv_and_cover_letter(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    cv_target = next(block for block in baseline["cv"]["blocks"] if block["type"] == "bullet")
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": [{
                    "material": "cv",
                    "baseline_id": cv_target["id"],
                    "action": "rewrite",
                    "text": "Reviewed vendor contracts and recorded actionable operations follow-up.",
                    "jd_anchor_ids": ["JD-001"],
                }],
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "blocked"
    assert "canonical_draft_invalid" in outcome["blockers"]
    assert "baseline_transform_material_missing:cover_letter" in outcome["error"]


def test_near_total_rewrite_is_rejected_but_focused_rewriting_remains_available(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"
    cv_blocks = baseline["cv"]["blocks"]
    cl_target = next(block for block in baseline["cover_letter"]["blocks"] if block["type"] == "bullet")
    changes = [
        {
            "material": "cv",
            "baseline_id": block["id"],
            "action": "rewrite",
            "text": f"JD-specific rewrite {index} that still represents the source block.",
            "jd_anchor_ids": ["JD-001"],
        }
        for index, block in enumerate(cv_blocks[:6], start=1)
    ]
    changes.append(
        {
            "material": "cover_letter",
            "baseline_id": cl_target["id"],
            "action": "rewrite",
            "text": "Focused JD-specific Cover Letter evidence.",
            "jd_anchor_ids": ["JD-001"],
        }
    )

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": changes,
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "blocked"
    assert "baseline_transform_too_broad:cv" in outcome["error"]


@pytest.mark.legacy
def test_broad_but_not_replacement_delta_routes_to_stronger_audit(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"
    cv_targets = [block for block in baseline["cv"]["blocks"] if block["type"] in {"paragraph", "bullet", "heading"}][:4]
    cl_target = next(block for block in baseline["cover_letter"]["blocks"] if block["type"] == "bullet")
    changes = [
        {
            "material": "cv",
            "baseline_id": block["id"],
            "action": "rewrite",
            "text": f"JD-aligned evidence-preserving rewrite {index} for the selected Paralegal role.",
            "jd_anchor_ids": ["JD-001"],
        }
        for index, block in enumerate(cv_targets, start=1)
    ] + [
        {
            "material": "cover_letter",
            "baseline_id": cl_target["id"],
            "action": "rewrite",
            "text": "Contract review: I bring evidence-led operational support to the selected Paralegal role.",
            "jd_anchor_ids": ["JD-001"],
        }
    ]

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": changes,
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "succeeded"
    assert outcome["audit_task_packet"]["requires_strong_auditor"] is True
    assert outcome["audit_task_packet"]["model_routing"]["preferred_tier"] == "strong"


def test_small_transform_preserves_every_unmentioned_baseline_block(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001"},
    )["task_packet"]
    baseline = packet["content_baseline"]
    cv_target = next(
        block for block in baseline["cv"]["blocks"]
        if block["text"] == "A measurable action and result."
    )
    cl_target = next(
        block for block in baseline["cover_letter"]["blocks"]
        if "Supported the relevant work." in block["text"]
    )
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )["status"] == "succeeded"

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": [
                    {
                        "material": "cv",
                        "baseline_id": cv_target["id"],
                        "action": "rewrite",
                        "text": "Reviewed vendor contracts and converted findings into an accurate operations checklist.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                    {
                        "material": "cover_letter",
                        "baseline_id": cl_target["id"],
                        "action": "rewrite",
                        "text": "Contract review: I translated vendor-contract findings into clear operational follow-up for a payments team.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                ],
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "succeeded"
    draft = load_canonical_draft(package)
    cv_text = "\n".join(block["text"] for block in draft["cv"]["blocks"])
    assert "Reviewed vendor contracts and converted findings" in cv_text
    assert "Summary prototype" in cv_text
    assert set(draft["baseline_dispositions"]) == {
        block["id"]
        for material in ("cv", "cover_letter")
        for block in baseline[material]["blocks"]
        if block.get("content_floor", True)
    }
    assert all(
        disposition["action"] != "omit"
        for disposition in draft["baseline_dispositions"].values()
    )

    audit_task = outcome["audit_task_packet"]
    assert audit_task["audit_mode"] == "bounded_tailoring_delta"
    # Delta-first review still receives the compact final CV/CL so it can
    # catch role/employer drift, grammar damage and fragments globally.
    assert "materials.cv.text" in audit_task["read_allowlist"]
    assert "materials.cover_letter.text" in audit_task["read_allowlist"]
    changes = audit_task["tailoring_delta"]["changes"]
    cv_change = next(item for item in changes if item["baseline_ids"] == [cv_target["id"]])
    assert cv_change["before"] == ["A measurable action and result."]
    assert cv_change["after"] == "Reviewed vendor contracts and converted findings into an accurate operations checklist."
    assert audit_task["tailoring_delta"]["retained_block_count"] > 0


def test_complete_looking_replacement_cannot_bypass_the_baseline(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001"},
    )["task_packet"]
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "model_plan": plan},
    )["status"] == "succeeded"

    sparse_replacement = canonical_fixture()
    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={"job_id": "C0-001", "canonical_draft": sparse_replacement},
    )

    assert outcome["status"] == "blocked"
    assert "canonical_draft_invalid" in outcome["blockers"]
    assert "baseline_transform_required" in outcome["error"]
    current = load_canonical_draft(package)
    assert current["baseline_sha256"] == packet["content_baseline"]["baseline_sha256"]
    assert "Summary prototype" in "\n".join(block["text"] for block in current["cv"]["blocks"])


@pytest.mark.legacy
def test_transform_can_reorder_existing_evidence_without_rewriting_the_cv(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    target = next(
        block for block in baseline["cv"]["blocks"]
        if block["text"] == "A measurable action and result."
    )
    cl_target = next(block for block in baseline["cover_letter"]["blocks"] if block["type"] == "bullet")
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": [
                    {
                        "material": "cv",
                        "baseline_id": target["id"],
                        "action": "reorder",
                        "position": 2,
                        "jd_anchor_ids": ["JD-001"],
                    },
                    {
                        "material": "cover_letter",
                        "baseline_id": cl_target["id"],
                        "action": "rewrite",
                        "text": "Contract review: I bring evidence-led vendor-contract support for operations teams.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                ],
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "succeeded"
    draft = load_canonical_draft(package)
    assert draft["cv"]["blocks"][2]["id"] == target["id"]
    assert draft["cv"]["blocks"][2]["text"] == "A measurable action and result."
    assert draft["baseline_dispositions"][target["id"]]["action"] == "reorder"


@pytest.mark.legacy
def test_transform_can_merge_baseline_blocks_only_with_traceable_coverage(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    source_ids = [
        next(block["id"] for block in baseline["cv"]["blocks"] if token in block["text"])
        for token in ("Summary prototype", "Core label - core evidence")
    ]
    cl_target = next(block for block in baseline["cover_letter"]["blocks"] if block["type"] == "bullet")
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"

    outcome = dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": [
                    {
                        "material": "cv",
                        "baseline_ids": source_ids,
                        "action": "merge",
                        "text": "Paralegal candidate combining concise role positioning with evidence-led contract operations expertise.",
                        "type": "paragraph",
                        "section": "summary",
                        "jd_anchor_ids": ["JD-001"],
                    },
                    {
                        "material": "cover_letter",
                        "baseline_id": cl_target["id"],
                        "action": "rewrite",
                        "text": "Contract review: I bring evidence-led vendor-contract support for operations teams.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                ],
                "additions": [],
            },
        },
    )

    assert outcome["status"] == "succeeded"
    draft = load_canonical_draft(package)
    merged = next(block for block in draft["cv"]["blocks"] if block.get("change_action") == "merge")
    assert merged["baseline_refs"] == source_ids
    assert all(
        draft["baseline_dispositions"][source_id]["target_id"] == merged["id"]
        for source_id in source_ids
    )


@pytest.mark.legacy
def test_content_audit_refuses_a_canonical_file_with_silent_baseline_loss(tmp_path):
    workspace = build_workspace(tmp_path)
    package = build_package(workspace, with_outbound=False)
    packet = dispatch("materials", workspace=workspace, payload={"job_id": "C0-001"})["task_packet"]
    baseline = packet["content_baseline"]
    target = next(block for block in baseline["cv"]["blocks"] if block["type"] == "bullet")
    cl_target = next(block for block in baseline["cover_letter"]["blocks"] if block["type"] == "bullet")
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "model_plan": plan}
    )["status"] == "succeeded"
    assert dispatch(
        "materials",
        workspace=workspace,
        payload={
            "job_id": "C0-001",
            "canonical_draft": {
                "schema_version": 1,
                "artifact_type": "jobsflow_baseline_transform",
                    "job_id": "C0-001",
                    "baseline_sha256": baseline["baseline_sha256"],
                    **_drafting_binding(package),
                "changes": [
                    {
                        "material": "cv",
                        "baseline_id": target["id"],
                        "action": "rewrite",
                        "text": "Reviewed vendor contracts and recorded actionable follow-up for operations.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                    {
                        "material": "cover_letter",
                        "baseline_id": cl_target["id"],
                        "action": "rewrite",
                        "text": "Contract review: I bring evidence-led vendor-contract support for operations teams.",
                        "jd_anchor_ids": ["JD-001"],
                    },
                ],
                "additions": [],
            },
        },
    )["status"] == "succeeded"

    corrupted = load_canonical_draft(package)
    lost_id = baseline["cv"]["blocks"][0]["id"]
    corrupted["cv"]["blocks"] = [
        block for block in corrupted["cv"]["blocks"]
        if lost_id not in (block.get("baseline_refs") or [])
    ]
    atomic_write_json(package / "materials_draft.canonical.json", corrupted)

    outcome = dispatch(
        "materials", workspace=workspace, payload={"job_id": "C0-001", "stage": "drafting"}
    )

    assert outcome["status"] == "blocked"
    assert "baseline_content_floor_invalid" in outcome["blockers"]
    assert lost_id in outcome["error"]
