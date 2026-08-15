"""Regressions for fresh merge, artifact currentness, cursor, rescan, DOCX/MD."""

from __future__ import annotations

import json
from pathlib import Path

from tools.io_utils import atomic_write_json, atomic_write_text
from tools.workflow.engine import dispatch
from tools.workflow.entity_state import load_entity_state
from tools.workflow.fresh_store import FileFreshStore, FreshSnapshot
from tools.workflow.id_allocation import is_assigned_job_id
from tools.workflow.package_validator import MaterialsPackageValidator
from tools.workflow.testing_packages import build_package, build_workspace, prepare_package_for_apply, write_minimal_pdf


def _scored_csv(ws: Path, rows: list[dict], name: str = "fresh_24h_2026-08-14_twopass_scored.csv") -> Path:
    path = ws / "02_Tracker" / name
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = ["岗位编号", "职位", "公司", "链接", "CareerOps分数", "评估状态"]
    lines = [",".join(fields)]
    for row in rows:
        lines.append(",".join(row.get(k, "") for k in fields))
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def _register_run(ws: Path, run_id: str, scored: Path) -> None:
    from tools.workflow.adapters.scan import write_run_record

    write_run_record(ws, run_id=run_id, mode="temp", scored_path=scored)


def test_push_merges_and_keeps_old_fresh_rows(tmp_path):
    ws = build_workspace(tmp_path)
    title = "fresh_24h_keep"
    store = FileFreshStore(
        ws,
        title,
        [{"岗位编号": "OLD-1", "职位": "OLD", "公司": "Keep", "链接": "https://example.test/job/old"}],
    )
    scored = _scored_csv(
        ws,
        [{"岗位编号": "NEW-1", "职位": "NEW", "公司": "Acme", "链接": "https://example.test/job/new", "CareerOps分数": "4.0", "评估状态": ""}],
    )
    _register_run(ws, "run-merge", scored)
    preview = dispatch("push", workspace=ws, store=store, payload={"run_id": "run-merge", "fresh_title": title})
    assert preview["status"] == "planned"
    out = dispatch(
        "push",
        workspace=ws,
        store=store,
        payload={
            "run_id": "run-merge",
            "fresh_title": title,
            "confirmation_id": preview["proposal_id"],
        },
    )
    assert out["status"] == "succeeded"
    ids = {(row.get("岗位编号") or "") for row in store.read_active().rows}
    titles = {(row.get("职位") or "") for row in store.read_active().rows}
    assert "OLD-1" in ids
    assert "OLD" in titles
    assert "NEW-1" not in ids
    assert any(is_assigned_job_id(job_id) for job_id in ids)


def test_editing_cv_or_plan_invalidates_apply_ready(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)
    first = dispatch("apply", workspace=ws, payload={"job_id": "C0-001"})
    assert first["apply_ready"] is True
    (package / "cv.txt").write_text("rewritten body without matching hashes\n", encoding="utf-8")
    after_cv = dispatch("apply", workspace=ws, payload={"job_id": "C0-001"})
    assert after_cv["apply_ready"] is False
    assert "stale_artifact" in after_cv["blockers"] or "stale_input_used" in after_cv["blockers"]

    package2 = build_package(ws, job_id="C0-002")
    prepare_package_for_apply(ws, job_id="C0-002")
    first2 = dispatch("apply", workspace=ws, payload={"job_id": "C0-002"})
    assert first2["apply_ready"] is True
    plan = json.loads((package2 / "materials_plan.validated.json").read_text(encoding="utf-8"))
    plan["themes"] = ["edited-after-audit"]
    (package2 / "materials_plan.validated.json").write_text(json.dumps(plan), encoding="utf-8")
    after_plan = dispatch("apply", workspace=ws, payload={"job_id": "C0-002"})
    assert after_plan["apply_ready"] is False


def test_score_failure_does_not_commit_refresh_cursor(tmp_path, monkeypatch):
    ws = build_workspace(tmp_path)
    tracker = ws / "02_Tracker"
    state_path = tracker / "fresh_refresh_state.json"
    atomic_write_json(
        state_path,
        {
            "version": 1,
            "last_refresh_at": "2026-08-01T00:00:00Z",
            "last_mode": "temp",
            "last_window_hours": 24,
            "history": [],
        },
    )
    from tools.workflow.adapters import scan as scan_mod

    def fake_run(cmd, **kwargs):
        class P:
            returncode = 1 if "two_pass_score" in " ".join(str(c) for c in cmd) else 0
            stderr = "score failed"
            stdout = ""

        if "fresh_24h_scan.py" in " ".join(str(c) for c in cmd):
            assert "--no-record" in cmd
            (tracker / "fresh_24h_2026-08-14.csv").write_text("url,title\n", encoding="utf-8")
            atomic_write_json(
                tracker / "fresh_24h_2026-08-14_run.json",
                {"mode": "temp", "hours": 3, "window": {"since": None}, "counts": {"new": 1}, "candidates_csv": "x"},
            )
        return P()

    monkeypatch.setattr(scan_mod.subprocess, "run", fake_run)
    before = json.loads(state_path.read_text(encoding="utf-8"))["last_refresh_at"]
    out = scan_mod.default_scan_runner({"mode": "temp", "run_id": "run-fail"}, ws)
    assert out["status"] == "failed"
    after = json.loads(state_path.read_text(encoding="utf-8"))["last_refresh_at"]
    assert after == before


def test_two_temp_scans_use_distinct_run_ids(tmp_path):
    ws = build_workspace(tmp_path)
    first = dispatch(
        "scan",
        workspace=ws,
        payload={"mode": "temp", "run_id": "scan-aaaa1111", "fixture": {"jobs": [{"job_id": "A0-001"}]}},
    )
    second = dispatch(
        "scan",
        workspace=ws,
        payload={"mode": "temp", "run_id": "scan-bbbb2222", "fixture": {"jobs": [{"job_id": "A0-002"}]}},
    )
    assert first["status"] == "succeeded"
    assert second["status"] == "succeeded"
    assert load_entity_state(ws, "scan", "scan-aaaa1111").phase == "scan_completed"
    assert load_entity_state(ws, "scan", "scan-bbbb2222").phase == "scan_completed"


def test_apply_accepts_docx_pdf_and_markdown_email(tmp_path):
    from docx import Document

    ws = build_workspace(tmp_path)
    package = build_package(ws, job_id="C0-010", with_outbound=False)
    for name in ("cv.txt", "cl.txt", "email.txt"):
        path = package / name
        if path.exists():
            path.unlink()
    doc = Document()
    doc.add_paragraph("Paralegal at Acme. Reviewed vendor contracts. IELTS 7.5.")
    doc.save(package / "Pat_CV_Acme.docx")
    doc2 = Document()
    doc2.add_paragraph("Paralegal at Acme: I can bring adjacent contract-review experience. IELTS 7.5.")
    doc2.save(package / "Pat_Cover_Letter_Acme.docx")
    write_minimal_pdf(package / "Pat_CV_Acme.pdf", "Paralegal at Acme. CV IELTS 7.5")
    write_minimal_pdf(package / "Pat_Cover_Letter_Acme.pdf", "Paralegal at Acme. CL IELTS 7.5")
    atomic_write_text(package / "application_email.md", "Application for Paralegal at Acme. IELTS 7.5.\n")
    # File-format compatibility is independent from the required semantic
    # child audit; a deterministic legacy audit can no longer mint a receipt.
    report = MaterialsPackageValidator().validate(package, require_audit_receipt=False)
    out = dispatch("apply", workspace=ws, payload={"job_id": "C0-010"})
    assert not any(item["code"] in {"missing_cv", "missing_cl", "missing_email"} for item in report["findings"])
    assert out["apply_ready"] is False
    assert "content_audit_missing" in out["blockers"]
    codes = {item["code"] for item in report["findings"]}
    assert "missing_cv" not in codes
    assert "missing_email" not in codes
