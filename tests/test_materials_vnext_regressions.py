"""Regression coverage for the fixed baseline/template hand-off."""

from __future__ import annotations

import json

from docx import Document

from tools.workflow.materials_renderer import (
    _add_block,
    _job_heading_parts,
    _template_prototypes,
    mechanical_format_gate,
    render_canonical_docx,
)
from tools.workflow.materials_vnext.baseline import compile_baseline
from tools.workflow.materials_vnext.bundle import bundle_path
from tools.workflow.materials_vnext.engine import MaterialsEngine
from tools.workflow.package_context import PackageContextLoader
from tools.workflow.testing_packages import build_package, build_workspace, prepare_package_for_apply


def _plan_and_bundle(ws, package):
    first = MaterialsEngine().handle({"job_id": "C0-001", "stage": "plan"}, workspace=ws)
    assert first["status"] == "succeeded"
    planned = MaterialsEngine().handle(
        {
            "job_id": "C0-001",
            "stage": "plan",
            "model_plan": {
                "task_type": "materials_plan_and_bounded_tailoring",
                "duties": ["review and manage compliance processes"],
                "themes": ["process management"],
                "match_type": "direct_or_transferable",
            },
        },
        workspace=ws,
    )
    assert planned["after_state"] == "plan_ready"
    return json.loads(bundle_path(package).read_text(encoding="utf-8"))


def test_baseline_marks_compact_education_and_qualifications_outside_experience(tmp_path):
    ws = build_workspace(tmp_path)
    master = ws / "01_Masters" / "C_track" / "master_C_test_v1.docx"
    document = Document(str(master))
    section = document.add_paragraph(style="Resume Section")
    section.add_run("EDUCATION")
    education = document.add_paragraph(style="Compact Line")
    education.add_run("LL.M., University of Hong Kong")
    section = document.add_paragraph(style="Resume Section")
    section.add_run("QUALIFICATIONS & LANGUAGES")
    qualifications = document.add_paragraph(style="Compact Line")
    qualifications.add_run("English: IELTS 7.5")
    document.save(master)

    baseline = compile_baseline(workspace=ws, lane="C", role="Compliance Officer", candidate_name="Test Candidate")
    education_block = next(item for item in baseline["cv"]["blocks"] if item["text"].startswith("LL.M."))
    qualification_block = next(item for item in baseline["cv"]["blocks"] if item["text"].startswith("English:"))
    assert education_block["experience_id"] == ""
    assert qualification_block["experience_id"] == ""
    assert education_block["presentation_role"] == "compact_line"
    assert qualification_block["presentation_role"] == "compact_line"
    assert education_block["source_style"] == "Compact Line"


def test_renderer_uses_master_compact_style_for_education_and_date_tabs(tmp_path):
    ws = build_workspace(tmp_path)
    master = ws / "01_Masters" / "C_track" / "master_C_test_v1.docx"
    template = Document(str(master))
    prototypes = _template_prototypes(template, material="cv")
    document = Document(str(master))
    # Clear fixture paragraphs while retaining styles and page setup.
    body = document._element.body
    for child in list(body):
        if child.tag.rsplit("}", 1)[-1] != "sectPr":
            body.remove(child)
    _add_block(
        document,
        {
            "id": "education",
            "type": "bullet",
            "text": "LL.M., University of Hong Kong",
            "section": "education",
            "presentation_role": "compact_line",
            "source_style": "Compact Line",
        },
        material="cv",
        position=0,
        prototypes=prototypes,
    )
    _add_block(
        document,
        {
            "id": "job",
            "type": "heading",
            "text": "Compliance Officer Jan 2022 - Present",
            "section": "experience",
            "experience_id": "experience-01",
            "presentation_role": "job_heading",
            "source_style": "Job Heading",
        },
        material="cv",
        position=1,
        prototypes=prototypes,
    )
    assert document.paragraphs[0].style.name == "Compact Line"
    assert document.paragraphs[1].style.name == "Job Heading"
    assert "\t" in document.paragraphs[1].text
    assert _job_heading_parts("Compliance Officer Jan 2022 - Present") == ("Compliance Officer", "Jan 2022 - Present")
    # Style names alone are insufficient: the regression that prompted this
    # test copied ``Education`` as a generic black/normal paragraph while the
    # lane master uses a two-run compact-line contract.  Compare the direct
    # OOXML run properties (font, colour, weight and size) as the format gate
    # does for every rendered block.
    def rpr_shape(element):
        def walk(node):
            local = str(node.tag).rsplit("}", 1)[-1]
            attrs = tuple(sorted((str(key).rsplit("}", 1)[-1], str(value)) for key, value in node.attrib.items()))
            return (local, attrs, str(node.text or ""), tuple(walk(child) for child in node))

        return walk(element) if element is not None else ()

    assert rpr_shape(document.paragraphs[0].runs[0]._r.rPr) == rpr_shape(prototypes["compact"]["rprs"][0])
    assert rpr_shape(document.paragraphs[0].runs[1]._r.rPr) == rpr_shape(prototypes["compact"]["rprs"][1])
    assert rpr_shape(document.paragraphs[1].runs[1]._r.rPr) == rpr_shape(prototypes["job_heading"]["rprs"][1])


def test_cover_letter_pillar_preserves_label_separator(tmp_path):
    ws = build_workspace(tmp_path)
    master = ws / "01_Masters" / "C_track" / "cl_master_C_test_v1.docx"
    template = Document(str(master))
    prototypes = _template_prototypes(template, material="cover_letter")
    document = Document(str(master))
    body = document._element.body
    for child in list(body):
        if child.tag.rsplit("}", 1)[-1] != "sectPr":
            body.remove(child)

    _add_block(
        document,
        {
            "id": "pillar",
            "type": "bullet",
            "text": "Contract review - translated findings",
            "section": "pillar",
            "presentation_role": "baseline_block",
            "source_style": "Letter Bullet",
        },
        material="cover_letter",
        position=0,
        prototypes=prototypes,
    )

    assert document.paragraphs[0].text == "Contract review - translated findings"
    assert [run.text for run in document.paragraphs[0].runs] == [
        "Contract review",
        " - translated findings",
    ]


def test_baseline_numeric_evidence_loss_is_blocked_before_child_audit(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    bundle = _plan_and_bundle(ws, package)
    target = next(item for item in bundle["baseline"]["cv"]["blocks"] if "IELTS" in item["text"])
    transform = {
        "schema_version": 1,
        "operations": [
            {
                "material": "cv",
                "action": "replace",
                "target_id": target["id"],
                "before_text": target["text"],
                "after_text": "Summary of supported compliance work.",
                "jd_anchor_ids": ["JD-001"],
            }
        ],
    }
    result = MaterialsEngine().handle({"job_id": "C0-001", "transform": transform}, workspace=ws)
    assert result["status"] == "blocked"
    assert "baseline_content_preservation" in result["blockers"]
    assert not (package / "materials_vnext" / "audit_task.json").exists()


def test_audit_task_contains_compact_tailoring_delta(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    bundle = _plan_and_bundle(ws, package)
    target = next(item for item in bundle["baseline"]["cv"]["blocks"] if not item.get("host_managed") and item["type"] in {"paragraph", "bullet"})
    transform = {
        "schema_version": 1,
        "operations": [{
            "material": "cv",
            "action": "replace",
            "target_id": target["id"],
            "before_text": target["text"],
            "after_text": target["text"] + " Supports the priority process requirement.",
            "jd_anchor_ids": ["JD-001"],
        }],
    }
    result = MaterialsEngine().handle({"job_id": "C0-001", "transform": transform}, workspace=ws)
    assert result["status"] == "succeeded"
    task = result["audit_task_packet"]
    assert "tailoring_delta" in task["read_allowlist"]
    assert task["tailoring_delta"]["changed_block_count"] == 1
    assert task["tailoring_delta"]["retained_block_count"] > 0


def test_manifest_and_snapshot_are_reconciled_to_bound_package(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws, with_outbound=False)
    manifest_path = package / "job_manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["lane"] = "F"
    manifest["tier"] = {"code": "2", "label": "二级", "source": "legacy"}
    manifest["paths"]["package_dir"] = str(ws / "01_Masters" / "F_track" / "二级" / package.name)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    snapshot = package / "job_snapshot.md"
    snapshot.write_text(snapshot.read_text(encoding="utf-8").replace("Lane: C", "Lane: F").replace("Tier: 核心", "Tier: 二级"), encoding="utf-8")

    context = PackageContextLoader(ws).load("C0-001")
    repaired = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert context.package == str(package)
    assert repaired["lane"] == "C"
    assert repaired["tier"]["code"] == "0"
    assert repaired["paths"]["package_dir"] == str(package)
    assert "Lane: C" in snapshot.read_text(encoding="utf-8")
    assert "Tier: 核心" in snapshot.read_text(encoding="utf-8")


def test_format_gate_rejects_docx_changed_after_render_receipt(tmp_path):
    ws = build_workspace(tmp_path)
    package = build_package(ws)
    prepare_package_for_apply(ws)
    names = json.loads((package / "materials_render_receipt.json").read_text(encoding="utf-8"))["filenames"]
    document = Document(str(package / names["cv_docx"]))
    document.add_paragraph("post-render mutation")
    document.save(package / names["cv_docx"])
    report = mechanical_format_gate(package, ws)
    assert report["status"] == "failed"
    assert any(item["code"] == "render_receipt_hash_mismatch" for item in report["findings"])
