#!/usr/bin/env python3
"""Headless DOCX → PDF for cross-industry application packages.

Fully background—no WPS, GUI or Accessibility clicks. LibreOffice
`soffice --headless` is the documented path. Spire.Doc is an explicit fallback
only when LibreOffice is unavailable.

Usage:
  python3 tools/fresh_24h/docx_to_pdf.py path/to/file.docx
  python3 tools/fresh_24h/docx_to_pdf.py --package-dir path/to/package
"""

from __future__ import annotations

import argparse
import hashlib
import json
import subprocess
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from tools.io_utils import atomic_write_json

def _conversion_stamp_path(pdf: Path) -> Path:
    return pdf.with_suffix(pdf.suffix + ".jobsflow.json")


def _source_hash(docx: Path) -> str:
    digest = hashlib.sha256()
    with Path(docx).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_conversion_stamp(docx: Path, pdf: Path, *, engine: str) -> None:
    atomic_write_json(
        _conversion_stamp_path(Path(pdf)),
        {
            "source_sha256": _source_hash(Path(docx)),
            "engine": (engine or "libreoffice").lower(),
            "policy": "jobsflow-one-page-v1",
        },
    )


def conversion_cache_hit(docx: Path, pdf: Path, *, engine: str) -> bool:
    stamp = _conversion_stamp_path(Path(pdf))
    if not Path(docx).exists() or not Path(pdf).exists() or not stamp.exists():
        return False
    try:
        value = json.loads(stamp.read_text(encoding="utf-8"))
        return (
            value.get("source_sha256") == _source_hash(Path(docx))
            and value.get("engine") == (engine or "libreoffice").lower()
            and value.get("policy") == "jobsflow-one-page-v1"
        )
    except (OSError, ValueError, TypeError):
        return False


def _require_fixed_material_render_chain(docx: Path) -> None:
    """Prevent a model from converting a plain-text package DOCX directly.

    Lane application packages must be rendered by ``tools.workflow`` from a
    canonical draft and a lane master.  Generic DOCX/PDF conversion remains
    available for unrelated documents and the master files themselves, but a
    package containing ``job_manifest.json`` is a product artifact boundary.
    """

    package = Path(docx).parent
    if not (package / "job_manifest.json").is_file():
        return
    receipt = package / "materials_render_receipt.json"
    try:
        value = json.loads(receipt.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError):
        value = {}
    if not (
        str(value.get("renderer_version") or "").startswith("canonical-template-docx-")
        and isinstance(value.get("template_sha256"), dict)
        and isinstance(value.get("template_paths"), dict)
    ):
        raise RuntimeError(
            "fixed_material_entry_required: package DOCX/PDF must be produced by "
            "python3 -m tools.workflow materials render/pdf from the lane master"
        )


def find_soffice() -> str | None:
    candidates = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/opt/homebrew/bin/soffice"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice.bin"),
    ]
    which = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    if which.returncode == 0 and which.stdout.strip():
        candidates.insert(0, Path(which.stdout.strip()))
    for c in candidates:
        if c and c.exists():
            return str(c)
    return None


def convert_libreoffice(docx: Path, pdf: Path) -> bool:
    soffice = find_soffice()
    if not soffice:
        return False
    outdir = pdf.parent
    # Every conversion gets an isolated profile.  A shared /tmp profile makes
    # parallel CV/CL or multi-job conversion serialize on LibreOffice's lock
    # and can silently attach to the wrong process.
    with tempfile.TemporaryDirectory(prefix="jobsflow-lo-") as profile_dir:
        env_profile = Path(profile_dir).resolve().as_uri()
        subprocess.check_call(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--norestore",
                f"-env:UserInstallation={env_profile}",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(outdir),
                str(docx),
            ],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
    produced = outdir / (docx.stem + ".pdf")
    if produced.exists() and produced != pdf:
        produced.replace(pdf)
    return pdf.exists() and pdf.stat().st_size > 1000


def densify_docx(src: Path, dst: Path, *, is_cv: bool) -> None:
    from docx import Document
    from docx.shared import Pt, Twips

    d = Document(str(src))
    margin_factor = 0.88 if is_cv else 0.90
    font_delta = -0.5 if is_cv else -0.2
    for sec in d.sections:
        sec.top_margin = Twips(int(sec.top_margin.twips * margin_factor))
        sec.bottom_margin = Twips(int(sec.bottom_margin.twips * margin_factor))
        sec.left_margin = Twips(int(sec.left_margin.twips * margin_factor))
        sec.right_margin = Twips(int(sec.right_margin.twips * margin_factor))
    for p in d.paragraphs:
        pf = p.paragraph_format
        try:
            if pf.space_after is not None:
                pf.space_after = Pt(max(0, (pf.space_after.pt or 0) * 0.30))
            if pf.space_before is not None:
                pf.space_before = Pt(max(0, (pf.space_before.pt or 0) * 0.30))
            pf.line_spacing = 1.0
        except Exception:
            pass
        for run in p.runs:
            if run.font.size:
                try:
                    run.font.size = Pt(max(8.5, run.font.size.pt + font_delta))
                except Exception:
                    pass
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    try:
                        if pf.space_after is not None:
                            pf.space_after = Pt(max(0, (pf.space_after.pt or 0) * 0.30))
                        if pf.space_before is not None:
                            pf.space_before = Pt(max(0, (pf.space_before.pt or 0) * 0.30))
                    except Exception:
                        pass
    dst.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(dst))


def _content_bbox(page) -> tuple[float, float, float, float] | None:
    blocks = page.get_text("blocks") or []
    xs0, ys0, xs1, ys1 = [], [], [], []
    for b in blocks:
        if len(b) < 5:
            continue
        text = (b[4] or "").strip() if len(b) > 4 else ""
        if not text:
            continue
        if "Evaluation" in text or "Spire.Doc" in text:
            continue
        xs0.append(b[0])
        ys0.append(b[1])
        xs1.append(b[2])
        ys1.append(b[3])
    try:
        for d in page.get_drawings() or []:
            r = d.get("rect")
            if r is None:
                continue
            xs0.append(r.x0)
            ys0.append(r.y0)
            xs1.append(r.x1)
            ys1.append(r.y1)
    except Exception:
        pass
    if not xs0:
        return None
    return (min(xs0), min(ys0), max(xs1), max(ys1))


def strip_eval_warning_one_page(src_pdf: Path, dst_pdf: Path) -> int:
    import fitz

    doc = fitz.open(str(src_pdf))
    for page in doc:
        for needle in (
            "Evaluation Warning: The document was created with Spire.Doc for Python.",
            "Evaluation Warning",
            "Spire.Doc for Python",
            "created with Spire",
        ):
            for rect in page.search_for(needle):
                r = fitz.Rect(rect.x0 - 2, rect.y0 - 2, rect.x1 + 2, rect.y1 + 2)
                page.add_redact_annot(r, fill=(1, 1, 1))
        page.apply_redactions()

    if doc.page_count != 1:
        count = doc.page_count
        doc.close()
        raise RuntimeError(
            f"PDF has {count} pages; adjust DOCX spacing/content instead of scaling or overlaying pages"
        )

    doc.save(str(dst_pdf))
    n = doc.page_count
    doc.close()
    return n


def convert_spire(docx: Path, pdf: Path) -> bool:
    try:
        from spire.doc import Document as SpireDoc, FileFormat
    except ImportError:
        return False
    is_cv = "CV" in docx.name and "Cover" not in docx.name
    dense = Path("/tmp") / f"dense_{docx.stem}.docx"
    raw = Path("/tmp") / f"raw_{docx.stem}.pdf"
    densify_docx(docx, dense, is_cv=is_cv)
    doc = SpireDoc()
    doc.LoadFromFile(str(dense))
    doc.SaveToFile(str(raw), FileFormat.PDF)
    doc.Close()
    strip_eval_warning_one_page(raw, pdf)
    return pdf.exists() and pdf.stat().st_size > 5000


def convert(
    docx: Path,
    pdf: Path | None = None,
    *,
    engine: str = "libreoffice",
    force: bool = False,
    sanitize_metadata: bool = True,
) -> Path:
    """Convert DOCX → PDF. engine: auto | libreoffice | spire."""
    docx = docx.resolve()
    pdf = (pdf or docx.with_suffix(".pdf")).resolve()
    engine = (engine or "libreoffice").lower()
    _require_fixed_material_render_chain(docx)
    if sanitize_metadata:
        # Core properties are not CV/CL content and are therefore safe to
        # normalize before export.  This removes template residue early while
        # the semantic audit remains bound to normalized document text.
        try:
            from tools.workflow.materials_metadata import sanitize_docx_metadata

            sanitize_docx_metadata(docx, title=docx.stem)
        except Exception as exc:
            raise RuntimeError(f"DOCX metadata normalization failed: {exc}") from exc

    if not force and conversion_cache_hit(docx, pdf, engine=engine):
        print(f"OK cached (source unchanged): {pdf}")
        return pdf

    if engine in ("auto", "libreoffice"):
        if convert_libreoffice(docx, pdf):
            write_conversion_stamp(docx, pdf, engine=engine)
            print(f"OK libreoffice (headless): {pdf}")
            return pdf
        if engine == "libreoffice":
            raise RuntimeError(
                "LibreOffice soffice not found. Install LibreOffice, or use --engine spire."
            )

    if engine in ("auto", "spire"):
        if convert_spire(docx, pdf):
            write_conversion_stamp(docx, pdf, engine=engine)
            print(f"OK spire (headless): {pdf}")
            return pdf
        if engine == "spire":
            raise RuntimeError("Spire.Doc not available (pip install spire.doc).")

    raise RuntimeError(f"No headless converter available for {docx}")


def convert_package_dir(
    d: Path,
    *,
    engine: str = "libreoffice",
    force: bool = False,
    sanitize_metadata: bool = True,
) -> None:
    files = sorted(d.glob("*CV.docx")) + sorted(d.glob("*Cover Letter.docx"))
    if not files:
        files = [p for p in sorted(d.glob("*.docx")) if not p.name.startswith("~$")]
    for f in files:
        convert(f, engine=engine, force=force, sanitize_metadata=sanitize_metadata)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(
        description="Headless DOCX→PDF (LibreOffice first; explicit Spire fallback). Never launches WPS."
    )
    ap.add_argument("docx", nargs="?", type=Path)
    ap.add_argument("--package-dir", type=Path)
    ap.add_argument(
        "--engine",
        choices=("auto", "libreoffice", "spire"),
        default="libreoffice",
        help="libreoffice: documented/default path; use auto or spire only as an explicit fallback.",
    )
    ap.add_argument(
        "--force",
        action="store_true",
        help="Rebuild even when DOCX content and conversion policy are unchanged.",
    )
    ap.add_argument(
        "--preserve-metadata",
        action="store_true",
        help="Do not normalize DOCX core properties before export (not recommended for outbound materials).",
    )
    args = ap.parse_args(argv)

    if args.package_dir:
        convert_package_dir(
            args.package_dir,
            engine=args.engine,
            force=args.force,
            sanitize_metadata=not args.preserve_metadata,
        )
        return 0
    if not args.docx:
        ap.error("docx or --package-dir required")
    convert(args.docx, engine=args.engine, force=args.force, sanitize_metadata=not args.preserve_metadata)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
