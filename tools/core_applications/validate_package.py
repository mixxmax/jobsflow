#!/usr/bin/env python3
"""Validate generated core job application packages."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Optional, Sequence

from docx import Document
from pypdf import PdfReader

# The validator is also invoked as a file (``python tools/.../validate_package.py``)
# in release checks, where Python otherwise puts only ``tools/core_applications``
# on ``sys.path``.
REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from tools.job_materials.role_titles import build_role_title_contract
from tools.job_materials.publisher import RECRUITER_TYPES


REQUIRED_FILES = (
    ("job_snapshot.md", "job_snapshot.md"),
    ("申请指南.md", "申请指南.md"),
    ("application_log.md", "application_log.md"),
    ("CV.pdf", "*CV.pdf"),
    ("CV.docx", "*CV.docx"),
    ("Cover_Letter.pdf", "*Cover*Letter*.pdf"),
    ("Cover_Letter.docx", "*Cover*Letter*.docx"),
)

BANNED_TERMS = (
    "results-driven",
    "proven track record",
    "leverage",
    "spearhead",
    "delve",
    "testament",
    "【",
    "】",
    "TBD",
    "TODO",
    "JD candidate",
    "3.40/4.0",
    "90% satisfactory",
)

REQUIRED_JOB_FIELDS = ("company", "role", "lane", "parent_dir", "folder_name")


class ManifestError(ValueError):
    """Raised when a manifest does not match the validator's input contract."""


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    for section in document.sections:
        for area in (section.header, section.footer):
            parts.extend(paragraph.text for paragraph in area.paragraphs)
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return _docx_text(path)
    if path.suffix.lower() == ".pdf":
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    return path.read_text(encoding="utf-8")


def _has_text(text: str, expected: str) -> bool:
    normalized_text = " ".join(text.split()).casefold()
    normalized_expected = " ".join(expected.split()).casefold()
    return normalized_expected in normalized_text


def _role_tokens_match(role: str, text: str) -> bool:
    """Fuzzy role check: at least two significant role tokens appear in text.

    Material roles may legitimately extend the manifest role (e.g. "Paralegal /
    Legal Assistant (Private Funds)" vs "Paralegal (Private Funds)"), so a
    strict substring check would produce false alarms.
    """
    role_tokens = [
        tok
        for tok in re.findall(r"[A-Za-z0-9\u3400-\u4dbf\u4e00-\u9fff]+", role)
        if len(tok) >= 3
    ]
    if len(role_tokens) < 2:
        return False
    text_tokens = set(re.findall(r"[A-Za-z0-9\u3400-\u4dbf\u4e00-\u9fff]+", text.casefold()))
    matched = sum(1 for tok in role_tokens if tok.casefold() in text_tokens)
    return matched >= 2


def _validate_text(path: Path, company: str, role: str) -> list[str]:
    try:
        text = _read_text(path)
    except Exception as error:
        return [f"{path.name}: could not open text-bearing file ({error})"]

    errors = []
    folded_text = text.casefold()
    for term in BANNED_TERMS:
        if term.casefold() in folded_text:
            errors.append(f"{path.name}: banned wording found: {term}")
    if not _has_text(text, company):
        errors.append(f"{path.name}: missing company name: {company}")
    if not _has_text(text, role) and not _role_tokens_match(role, text):
        errors.append(f"{path.name}: missing role name: {role}")
    return errors


def _validate_pdf(path: Path, *, max_pages: int | None = None) -> list[str]:
    try:
        reader = PdfReader(path)
        has_valid_page = any(
            float(page.mediabox.width) > 0 and float(page.mediabox.height) > 0
            for page in reader.pages
        )
    except Exception as error:
        return [f"{path.name}: could not open PDF ({error})"]

    errors = []
    if not has_valid_page:
        errors.append(f"{path.name}: PDF has no page with nonzero dimensions")
    if max_pages is not None and len(reader.pages) > max_pages:
        errors.append(
            f"{path.name}: PDF has {len(reader.pages)} pages; maximum allowed is {max_pages}"
        )
    return errors


def _required_paths(package_dir: Path) -> tuple[list[Path], list[str]]:
    paths = []
    errors = []
    for label, pattern in REQUIRED_FILES:
        matches = sorted(package_dir.glob(pattern))
        if not matches:
            errors.append(f"missing required file: {label}")
        else:
            paths.append(matches[0])
    return paths, errors


_RESIDUAL_SENTENCE_RE = re.compile(
    # Do not flag valid list punctuation such as ``litigation support,
    # investigations and compliance``.  Only a verb left with no object (or
    # a comma stranded at the end of a line) is a residual fragment.
    r"\b(?:support|assist|coordinate|manage)\s+\.(?=\s|$)|"
    r"\b(?:support|assist|coordinate|manage)\s*,\s*(?=$|\n)",
    re.IGNORECASE,
)
_EMAIL_TEMPLATE_RE = re.compile(
    r"The role's focus on\s+my\s+[^.\n]{1,240}\s+"
    r"(?:matches|aligns with)\s+my\b",
    re.IGNORECASE,
)
_CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")


def _manifest_outbound_paths(package_dir: Path) -> list[Path]:
    """Return only files that can leave the private package."""
    patterns = (
        "*_CV.*",
        "*CV.*",
        "*_Cover_Letter.*",
        "*Cover*Letter*.*",
        "*cover*letter*.*",
        "*application*email*.*",
        "*email*.*",
    )
    paths: list[Path] = []
    for pattern in patterns:
        for path in sorted(package_dir.glob(pattern)):
            if not path.is_file() or path.name.endswith(".jobsflow.json"):
                continue
            if path not in paths:
                paths.append(path)
    return paths


def _normalized_token(value: str) -> str:
    return re.sub(r"[^a-z0-9\u3400-\u4dbf\u4e00-\u9fff]+", "", str(value or "").casefold())


def validate_manifest_contract(
    package_dir: Path,
    manifest: dict[str, Any],
    *,
    outbound_paths: list[Path] | None = None,
) -> list[str]:
    """Validate generated material against the per-job manifest contract."""
    package_dir = Path(package_dir).resolve()
    errors: list[str] = []
    if manifest.get("schema_version") != 1:
        errors.append("job_manifest.json: unsupported or missing schema_version")

    job = manifest.get("job") if isinstance(manifest.get("job"), dict) else {}
    job_id = str(manifest.get("job_id") or "").strip().upper()
    role = str(job.get("role_material") or job.get("role_display") or "").strip()
    role_display = str(job.get("role_display") or role).strip()
    publisher = str(job.get("publisher_name") or "").strip()
    publisher_type = str(job.get("publisher_type") or "unknown").strip().casefold()
    company_out = str(job.get("company_out") or "").strip()
    lane = str(manifest.get("lane") or "").strip().upper()
    tier = manifest.get("tier") if isinstance(manifest.get("tier"), dict) else {}
    if not job_id:
        errors.append("job_manifest.json: missing job_id")
    if not lane:
        errors.append("job_manifest.json: missing lane")
    if not role:
        errors.append("job_manifest.json: missing job.role_material")
    if not tier.get("code"):
        errors.append("job_manifest.json: missing tier.code")

    # New manifests carry a deterministic title contract.  Keep this check
    # tolerant of older packages, but reject a partially updated manifest that
    # would drop a substantive parenthetical or silently combine A/B roles.
    title_contract = job.get("role_title_contract")
    if isinstance(title_contract, dict):
        overrides = manifest.get("overrides") if isinstance(manifest.get("overrides"), dict) else {}
        expected_title = build_role_title_contract(
            role_display,
            selected_primary=str(overrides.get("role_primary") or ""),
        )
        if str(title_contract.get("primary") or "") != str(expected_title.get("primary") or ""):
            errors.append("job_manifest.json: role_title_contract.primary disagrees with role_display/override")
        if list(title_contract.get("alternates") or []) != list(expected_title.get("alternates") or []):
            errors.append("job_manifest.json: role_title_contract.alternates is stale")
        if role != str(expected_title.get("primary") or ""):
            errors.append("job_manifest.json: role_material is not the selected primary role")
        for parenthetical in expected_title.get("primary_parentheticals") or []:
            if parenthetical.get("kind") == "specialization":
                value = str(parenthetical.get("text") or "").strip()
                if value and f"({value})" not in role and f"（{value}）" not in role:
                    errors.append(
                        "job_manifest.json: substantive parenthetical was removed from role_material"
                    )

    paths = manifest.get("paths") if isinstance(manifest.get("paths"), dict) else {}
    recorded_package = str(paths.get("package_dir") or "").strip()
    if recorded_package:
        recorded_path = Path(recorded_package).expanduser()
        if not recorded_path.is_absolute():
            # Relative package paths are resolved from the package's parent so
            # manifests remain usable when a private workspace is relocated.
            recorded_path = package_dir.parent / recorded_path
        try:
            same_path = recorded_path.resolve() == package_dir
        except OSError:
            same_path = False
        if not same_path:
            errors.append(
                f"job_manifest.json: package path mismatch ({recorded_package} != {package_dir})"
            )
    if bool(paths.get("path_tier_mismatch")) or bool(
        (manifest.get("validation") or {}).get("path_tier_mismatch")
    ):
        errors.append("job_manifest.json: package tier does not match the job-id tier")

    expected_tier_label = str(tier.get("label") or "").strip()
    if expected_tier_label and expected_tier_label in {"核心", "一级", "二级"}:
        actual_tier_labels = {
            part
            for part in package_dir.parts
            if part in {"核心", "一级", "二级"}
        }
        if actual_tier_labels and expected_tier_label not in actual_tier_labels:
            errors.append(
                f"job_manifest.json: package directory tier is not {expected_tier_label}"
            )

    id_match = re.match(r"^[A-G][0-2]-", job_id)
    if id_match and str(tier.get("code")) != job_id[1]:
        errors.append(
            f"job_manifest.json: tier.code={tier.get('code')} disagrees with {job_id}"
        )

    outbound = outbound_paths if outbound_paths is not None else _manifest_outbound_paths(package_dir)
    if not outbound:
        errors.append("job_manifest.json: no outbound CV/Cover Letter/email files found")
        return errors

    validation = manifest.get("validation") if isinstance(manifest.get("validation"), dict) else {}
    language = str(
        validation.get("material_language")
        or (manifest.get("outbound") or {}).get("material_language")
        or "en"
    ).casefold()
    max_pages = validation.get("max_cover_letter_pages", 1)
    try:
        max_pages = int(max_pages)
    except (TypeError, ValueError):
        max_pages = 1
    publisher_token = _normalized_token(publisher)
    company_token = _normalized_token(company_out)
    for path in outbound:
        name_token = _normalized_token(path.name)
        if publisher_type in RECRUITER_TYPES and publisher_token and publisher_token in name_token:
            errors.append(f"{path.name}: recruiter/agency name leaked into outbound filename")
        if path.suffix.lower() == ".pdf":
            page_limit = max_pages if "cover" in path.name.casefold() else None
            errors.extend(_validate_pdf(path, max_pages=page_limit))
        try:
            text = _read_text(path)
        except Exception as error:
            errors.append(f"{path.name}: could not read outbound material ({error})")
            continue
        if text:
            if publisher_type in RECRUITER_TYPES and publisher_token:
                if publisher_token in _normalized_token(text):
                    errors.append(f"{path.name}: recruiter/agency name leaked into outbound text")
            if role and _normalized_token(role) not in _normalized_token(text):
                # The material role may extend the manifest role; fall back to
                # a fuzzy token check (any two significant role tokens).
                if not _role_tokens_match(role, text):
                    errors.append(
                        f"{path.name}: role name missing from outbound material: {role}"
                    )
            if company_token and company_token not in _normalized_token(text):
                errors.append(f"{path.name}: verified employer name missing from outbound material: {company_out}")
            for match in _RESIDUAL_SENTENCE_RE.finditer(text):
                errors.append(f"{path.name}: residual incomplete sentence near {match.group(0)!r}")
                break
            if path.name.casefold() == "application_email.txt" and (
                "[jd anchor:" in text.casefold() or _EMAIL_TEMPLATE_RE.search(text)
            ):
                errors.append(f"{path.name}: legacy JD-anchor sentence remains")
            if language.startswith("en") and _CJK_RE.search(text):
                errors.append(f"{path.name}: Chinese characters found in English outbound material")

    if language.startswith("en") and _CJK_RE.search(role):
        errors.append("job_manifest.json: English material role contains Chinese characters")
    return list(dict.fromkeys(errors))


def validate_package(
    package_dir: Path,
    company: str,
    role: str,
    *,
    job_manifest: dict[str, Any] | None = None,
) -> list[str]:
    """Return package contract violations; an empty list means validation passed.

    DOCX and one-page PDF are the product contract.  A legacy ``.tex`` source,
    when present, is still scanned for banned wording but is never required.
    """
    package_dir = Path(package_dir)
    if job_manifest is None:
        manifest_path = package_dir / "job_manifest.json"
        try:
            value = json.loads(manifest_path.read_text(encoding="utf-8"))
        except (OSError, ValueError, TypeError):
            value = None
        if isinstance(value, dict):
            job_manifest = value
    paths, errors = _required_paths(package_dir)
    for path in paths:
        if path.suffix.lower() == ".pdf":
            page_limit = None
            if job_manifest and "cover" in path.name.casefold():
                validation = job_manifest.get("validation")
                if isinstance(validation, dict):
                    try:
                        page_limit = int(validation.get("max_cover_letter_pages", 1))
                    except (TypeError, ValueError):
                        page_limit = 1
            errors.extend(_validate_pdf(path, max_pages=page_limit))
        else:
            # A recruiter-facing snapshot may correctly contain the publisher
            # rather than the verified employer.  Manifest validation checks
            # the outbound files separately, so do not apply the old employer
            # requirement to every internal text file in this mode.
            errors.extend(_validate_text(path, "" if job_manifest else company, role))
    # Keep legacy source files safe without making a LaTeX toolchain a release
    # prerequisite.  This also catches stale wording when a user keeps an old
    # source beside the current DOCX/PDF materials.
    for path in sorted(package_dir.glob("*.tex")):
        if path not in paths:
            errors.extend(_validate_text(path, "" if job_manifest else company, role))
    if job_manifest:
        errors.extend(validate_manifest_contract(package_dir, job_manifest))
    return list(dict.fromkeys(errors))


def _validate_manifest(manifest: object) -> list[dict]:
    if not isinstance(manifest, dict):
        raise ManifestError("top-level value must be an object containing a viable list")

    viable = manifest.get("viable")
    if not isinstance(viable, list):
        raise ManifestError("viable must be a list")

    for index, job in enumerate(viable):
        if not isinstance(job, dict):
            raise ManifestError(f"viable[{index}] must be an object")
        missing = [field for field in REQUIRED_JOB_FIELDS if field not in job]
        if missing:
            raise ManifestError(
                f"viable[{index}] missing required fields: {', '.join(missing)}"
            )
        invalid = [
            field
            for field in REQUIRED_JOB_FIELDS
            if not isinstance(job[field], str) or not job[field].strip()
        ]
        if invalid:
            raise ManifestError(
                f"viable[{index}] fields must be non-empty strings: {', '.join(invalid)}"
            )
    return viable


def _parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--manifest", type=Path, required=True)
    selection = parser.add_mutually_exclusive_group(required=True)
    selection.add_argument("--lane")
    selection.add_argument("--all", action="store_true", dest="all_jobs")
    return parser.parse_args(argv)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = _parse_args(argv)
    try:
        manifest = json.loads(args.manifest.read_text(encoding="utf-8"))
        jobs = _validate_manifest(manifest)
    except (OSError, json.JSONDecodeError, ManifestError) as error:
        print(f"manifest error: {error}")
        return 2

    selected_jobs = jobs if args.all_jobs else [job for job in jobs if job["lane"] == args.lane]
    if not selected_jobs:
        if args.all_jobs:
            print("manifest error: no packages selected for --all")
        else:
            available_lanes = ", ".join(sorted({job["lane"] for job in jobs})) or "none"
            print(
                f"manifest error: no packages selected for lane {args.lane!r}; "
                f"available lanes: {available_lanes}"
            )
        return 2

    failed = 0
    for job in selected_jobs:
        package_dir = Path(job["parent_dir"]) / job["folder_name"]
        errors = validate_package(package_dir, job["company"], job["role"])
        if errors:
            failed += 1
            print(f"FAIL {job['company']} - {job['role']} ({package_dir})")
            for error in errors:
                print(f"  - {error}")
        else:
            print(f"PASS {job['company']} - {job['role']} ({package_dir})")

    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
