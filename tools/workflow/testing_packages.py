"""Synthetic JobSearch-shaped packages. No real candidate data."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path

from tools.io_utils import atomic_write_json, atomic_write_text
from tools.job_materials.jd_store import read_jd
from tools.workflow.plan_gate import write_validated_plan

JD = (
    "Key Responsibilities\n"
    "Draft and review vendor contracts for the operations team. " * 12
    + "\nRequirements\nLaw degree and three years of contract experience. " * 8
)


def _sha(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def write_minimal_pdf(path: Path, text: str = "Hello", pages: int = 1) -> None:
    # Tiny but valid PDF with a text layer.  Include xref offsets so the
    # production PDF validator and pypdf agree that this is a readable file.
    objects: list[str] = ["<< /Type /Catalog /Pages 2 0 R >>"]
    kids = " ".join(f"{3 + i * 2} 0 R" for i in range(pages))
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {pages} >>")
    for _ in range(pages):
        objects.extend(["", ""])
    font_id = len(objects) + 1
    objects.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for index in range(pages):
        page_id = 3 + index * 2
        content_id = page_id + 1
        safe_text = str(text or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET"
        objects[page_id - 1] = (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Contents {content_id} 0 R /Resources << /Font << /F1 {font_id} 0 R >> >> >>"
        )
        objects[content_id - 1] = f"<< /Length {len(stream.encode('latin-1'))} >>\nstream\n{stream}\nendstream"
    body = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body.extend(f"{number} 0 obj\n{obj}\nendobj\n".encode("latin-1"))
    xref = len(body)
    body.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    body.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        body.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    body.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(bytes(body))


def build_workspace(root: Path) -> Path:
    ws = root / "JobSearch_2026"
    (ws / "00_Profile").mkdir(parents=True)
    (ws / "01_Masters" / "C_track" / "核心").mkdir(parents=True)
    (ws / "02_Tracker" / "job_assessments").mkdir(parents=True)
    atomic_write_json(ws / "00_Profile" / "queries.json", {"schema_version": 2, "setup_required": False})
    # Materials rendering must never silently fall back to a generic
    # ``Candidate`` filename.  Keep the synthetic fixture explicit so tests
    # exercise the same confirmed-name gate as a real runtime.
    atomic_write_json(ws / "00_Profile" / "config.personal.json", {"candidate_name": "Test Candidate"})
    atomic_write_json(
        ws / "00_Profile" / "fact_evidence.json",
        {
            "nodes": [
                {"id": "EVID-AAA", "text": "Reviewed vendor contracts for a payments team.", "source": "base"},
                {"id": "EVID-BBB", "text": "Prepared bilingual closing checklists.", "source": "base"},
            ],
            "forbidden_claims": ["admitted as a solicitor"],
        },
    )
    _write_test_lane_templates(ws)
    return ws


def _write_test_lane_templates(ws: Path) -> None:
    """Create explicit lane masters for synthetic workspaces.

    Tests must exercise the same invariant as a real workspace: the renderer
    receives a lane master and may not silently fall back to a blank document.
    These are deliberately small style-bearing DOCX fixtures, not product
    templates or candidate data.
    """

    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    lane = ws / "01_Masters" / "C_track"
    cv_path = lane / "master_C_test_v1.docx"
    cl_path = lane / "cl_master_C_test_v1.docx"

    def add_style(document, name: str, size: float, *, bold: bool = False) -> None:
        styles = document.styles
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.0

    cv = Document()
    add_style(cv, "Resume Section", 11, bold=True)
    add_style(cv, "Job Heading", 10, bold=True)
    add_style(cv, "Resume Bullet", 9.5)
    add_style(cv, "Compact Line", 9.5)

    def format_run(run, *, size: float, bold: bool = False, color: str = "1C1C1C") -> None:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)

    title = cv.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(title.add_run("Test Candidate"), size=20.5, bold=True, color="17365D")
    contact = cv.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(contact.add_run("Hong Kong | test@example.test"), size=9, bold=True, color="4B4B4B")
    summary = cv.add_paragraph()
    format_run(summary.add_run("Summary prototype"), size=9.5)
    section = cv.add_paragraph(style="Resume Section")
    format_run(section.add_run("PROFESSIONAL SUMMARY"), size=11, bold=True, color="17365D")
    compact = cv.add_paragraph(style="Compact Line")
    format_run(compact.add_run("Core label"), size=8.5, bold=True, color="17365D")
    format_run(compact.add_run(" - core evidence"), size=8.5)
    heading = cv.add_paragraph(style="Job Heading")
    tabs = heading.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.0))
    format_run(heading.add_run("Role | Company"), size=9.5, bold=True)
    format_run(heading.add_run("\t2026"), size=9.5, bold=True, color="4B4B4B")
    bullet = cv.add_paragraph(style="Resume Bullet")
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    bullet._p.get_or_add_pPr().append(num_pr)
    format_run(bullet.add_run("A measurable action and result."), size=9.5)
    cv.save(cv_path)

    cl = Document()
    add_style(cl, "Letter Body", 11)
    add_style(cl, "Letter Bullet", 10.5)
    add_style(cl, "Letter Compact", 10.5)
    title = cl.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(title.add_run("Test Candidate"), size=20.5, bold=True, color="17365D")
    contact = cl.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(contact.add_run("Hong Kong | test@example.test"), size=9, bold=True, color="4B4B4B")
    compact = cl.add_paragraph(style="Letter Compact")
    format_run(compact.add_run("[Date]"), size=10.5)
    subject = cl.add_paragraph()
    format_run(subject.add_run("Re: Application for [Role]"), size=11, bold=True, color="17365D")
    body = cl.add_paragraph(style="Letter Body")
    format_run(body.add_run("Dear Hiring Manager,"), size=11)
    letter_bullet = cl.add_paragraph(style="Letter Bullet")
    format_run(letter_bullet.add_run("Evidence: "), size=10.5, bold=True, color="17365D")
    format_run(letter_bullet.add_run("Supported the relevant work."), size=10.5)
    cl.save(cl_path)


def prepare_package_for_apply(ws: Path, job_id: str = "C0-001") -> None:
    """Drive a synthetic package through the public workflow gates."""
    from tools.workflow.engine import dispatch
    from tools.workflow.materials_renderer import expected_filenames

    package = ws / "01_Masters" / "C_track" / "核心" / f"{job_id}_未投_Acme"
    plan = json.loads((package / "materials_plan.validated.json").read_text(encoding="utf-8"))
    assert dispatch("materials", workspace=ws, payload={"job_id": job_id, "model_plan": plan})["status"] == "succeeded"
    drafted = dispatch(
        "materials",
        workspace=ws,
        payload={"job_id": job_id, "canonical_draft": canonical_fixture(job_id)},
    )
    assert drafted["status"] == "succeeded"
    task = drafted["audit_task_packet"]
    passed = {
        "job_id": job_id,
        "audit_scope": "jd_mapping_and_presentation",
        "audit_input_fingerprint": task["audit_input_fingerprint"],
        "auditor_context_id": task["auditor_context_id"],
        "counts": {"P0": 0, "P1": 0, "P2": 0},
        "findings": [],
    }
    assert dispatch("audit", workspace=ws, payload={"job_id": job_id, "audit_result": passed})["status"] == "succeeded"
    assert dispatch("materials", workspace=ws, payload={"job_id": job_id, "stage": "render"})["status"] == "succeeded"
    names = expected_filenames(package, ws)
    cv_text = "Paralegal contract review operations experience with accurate records and reliable stakeholder support."
    cl_text = "Application for Paralegal at Acme with evidence based contract review support and operational value."
    write_minimal_pdf(package / names["cv_pdf"], cv_text)
    write_minimal_pdf(package / names["cl_pdf"], cl_text)
    assert dispatch("materials", workspace=ws, payload={"job_id": job_id, "stage": "pdf_generated"})["status"] == "succeeded"
    assert dispatch("format", workspace=ws, payload={"job_id": job_id})["status"] == "succeeded"


def canonical_fixture(job_id: str = "C0-001") -> dict:
    """Complete structured text used by synthetic workflow tests."""

    return {
        "schema_version": 1,
        "artifact_type": "jobsflow_canonical_cv_cl",
        "job_id": job_id,
        "cv": {
            "blocks": [
                {"id": "cv-heading", "type": "heading", "text": "Paralegal", "claim_ids": []},
                {
                    "id": "cv-summary",
                    "type": "paragraph",
                    "text": "Paralegal candidate bringing careful contract review support, reliable records management, bilingual coordination, IELTS 7.5 English and evidence based operational follow through for busy legal and business teams.",
                    "claim_ids": ["C1"],
                },
                {
                    "id": "cv-evidence-1",
                    "type": "bullet",
                    "text": "Reviewed vendor contracts in an adjacent payments setting and maintained clear checklists for accurate stakeholder follow up.",
                    "claim_ids": ["C1"],
                },
            ]
        },
        "cover_letter": {
            "blocks": [
                {
                    "id": "cl-opening",
                    "type": "paragraph",
                    "text": "I am applying for the Paralegal role at Acme. Its focus on vendor contract review is relevant to my adjacent payments experience and disciplined document support.",
                    "claim_ids": ["C1"],
                },
                {
                    "id": "cl-evidence-1",
                    "type": "paragraph",
                    "text": "I have reviewed vendor contracts and maintained practical checklists, and my IELTS 7.5 English supports accurate records, responsive coordination and dependable operational support without overstating ownership.",
                    "claim_ids": ["C1"],
                },
                {"id": "cl-signoff", "type": "signoff", "text": "Yours sincerely,\nCandidate", "claim_ids": []},
            ]
        },
    }


def build_package(
    ws: Path,
    job_id: str = "C0-001",
    *,
    full_jd: bool = True,
    assessment_stale: bool = False,
    unanswered: bool = False,
    publisher_type: str = "recruiter",
    publisher_name: str = "Michael Page",
    with_plan: bool = True,
    with_outbound: bool = True,
    recruiter_in_name: bool = False,
    transferable_as_direct: bool = False,
    language_mismatch: bool = False,
    missing_attachment: bool = False,
    pdf_pages: int = 1,
    pdf_text: bool = True,
) -> Path:
    package = ws / "01_Masters" / "C_track" / "核心" / f"{job_id}_未投_Acme"
    package.mkdir(parents=True, exist_ok=True)
    jd = JD if full_jd else "Short teaser only."
    atomic_write_text(
        package / "jd_full.md",
        f"# JD — {job_id}\n\n- source: user_paste\n\n---\n\n{jd}\n",
    )
    atomic_write_text(package / "job_snapshot.md", f"Role: Paralegal\nCompany: Acme\nURL: https://example.test/{job_id}\n")
    jd_hash = _sha(read_jd(package, ws))
    atomic_write_json(
        package / "assessment.json",
        {
            "job_id": job_id,
            "match_type": "transferable",
            "revision": 1,
            "jd_hash": "stale" if assessment_stale else jd_hash,
            "strengths": [{"text": "contract review", "evidence_id": "EVID-AAA"}],
            "gaps": [],
        },
    )
    atomic_write_json(
        package / "application_preflight.json",
        {
            "ready_for_apply": not unanswered,
            "unanswered_hard": ["work_auth"] if unanswered else [],
            "questions": ["work_auth"] if unanswered else [],
        },
    )
    atomic_write_json(
        package / "job_manifest.json",
        {
            "schema_version": 1,
            "job_id": job_id,
            "lane": job_id[:1],
            "tier": {"code": job_id[1:2] or "0", "label": "核心"},
            "job": {
                "role_material": "Paralegal",
                "role_display": "Paralegal",
                "publisher_type": publisher_type,
                "publisher_name": publisher_name,
                "company_out": "Acme",
            },
        },
    )
    claim_text = (
        "This maps directly to the core duty and is direct experience."
        if transferable_as_direct
        else "Reviewed vendor contracts in an adjacent payments setting."
    )
    plan = {
        "task_type": "materials_plan",
        "duties": ["Draft vendor contracts"],
        "themes": ["contracts"],
        "match_type": "transferable",
        "claim_ledger": [
            {
                "id": "C1",
                "text": claim_text,
                "evidence_id": "EVID-AAA",
                "kind": "Direct" if transferable_as_direct else "Transferable",
                "assessment": "transferable",
            }
        ],
        "gaps": [],
        "differentiation": "contract operations",
    }
    if with_plan:
        write_validated_plan(package, plan)
    (package / "attachments").mkdir(exist_ok=True)
    if not missing_attachment:
        (package / "attachments" / "degree.pdf").write_bytes(b"%PDF-1.1\n%%EOF\n")
    cv_name = f"Pat_Paralegal_{'Michael_Page' if recruiter_in_name else 'Acme'}.pdf"
    cl_level = "IELTS 8.0" if language_mismatch else "IELTS 7.5"
    if with_outbound:
        atomic_write_text(
            package / "cv.txt",
            f"Paralegal at Acme. Reviewed vendor contracts. IELTS 7.5. {claim_text}",
        )
        atomic_write_text(
            package / "cl.txt",
            f"Paralegal at Acme: I can bring adjacent contract-review experience. {cl_level}.",
        )
        atomic_write_text(
            package / "email.txt",
            "Application for Paralegal at Acme. Please find my materials. IELTS 7.5.",
        )
        outbound_company = "Michael_Page" if recruiter_in_name else "Acme"
        write_minimal_pdf(
            package / f"Pat_Paralegal_{outbound_company}_CV.pdf",
            "Paralegal at Acme. CV IELTS 7.5" if pdf_text else "",
            pages=pdf_pages,
        )
        write_minimal_pdf(
            package / f"Pat_Paralegal_{outbound_company}_Cover_Letter.pdf",
            "Paralegal at Acme. CL IELTS 7.5" if pdf_text else "",
            pages=1,
        )
    packet = {
        "task_type": "materials_plan",
        "full_jd": full_jd,
        "facts": True,
        "assessment": {"match_type": "transferable"},
        "preflight": {"unanswered_hard": []},
        "evidence_ids": ["EVID-AAA", "EVID-BBB"],
        "publisher_type": publisher_type,
        "publisher_name": publisher_name,
        "input_hashes": {"jd": jd_hash},
        "stale_reasons": ["assessment_jd_hash"] if assessment_stale else [],
        "claim_ledger": plan["claim_ledger"],
        "outbound": {
            "required_attachments": ["degree.pdf"],
            "language_levels": {
                "cv": "IELTS 7.5",
                "cl": cl_level,
                "email": "IELTS 7.5",
            },
            "numbers": {"cv": ["7.5"], "cl": ["8.0" if language_mismatch else "7.5"], "email": ["7.5"]},
        },
    }
    from tools.workflow.package_context import PackageContextLoader

    ctx = PackageContextLoader(ws).load(job_id)
    if ctx.input_hashes and not assessment_stale:
        packet["input_hashes"] = dict(ctx.input_hashes)
    atomic_write_json(package / "materials_task_packet.json", packet)
    if with_plan and with_outbound and not assessment_stale and not unanswered and not missing_attachment and not language_mismatch and not transferable_as_direct and not recruiter_in_name:
        from tools.workflow.package_validator import audit_package

        audit_package(package)
    return package
