"""Semantic, metadata and container hashes for materials lifecycle gates."""

from __future__ import annotations

import hashlib
import json
import re
import zipfile
from pathlib import Path
from typing import Any


def normalize_text(text: str) -> str:
    lines = []
    for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        value = re.sub(r"[ \t]+", " ", line).strip()
        if value:
            lines.append(value)
    return "\n".join(lines)


def _sha_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_text(text: str) -> str:
    return _sha_bytes(str(text or "").encode("utf-8"))


def read_material_text(path: Path) -> str:
    path = Path(path)
    if not path.is_file():
        return ""
    if path.suffix.casefold() in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if path.suffix.casefold() == ".docx":
        try:
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(paragraphs)
        except Exception:
            return ""
    if path.suffix.casefold() == ".pdf":
        try:
            from pypdf import PdfReader

            return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
        except Exception:
            return ""
    return ""


def semantic_hash(path: Path) -> str:
    return sha256_text(normalize_text(read_material_text(Path(path))))


def container_hash(path: Path) -> str:
    try:
        return _sha_bytes(Path(path).read_bytes())
    except OSError:
        return ""


def metadata_hash(path: Path) -> str:
    """Hash only metadata, so metadata-only edits don't force semantic review."""

    path = Path(path)
    if not path.is_file():
        return ""
    metadata: dict[str, Any] = {}
    if path.suffix.casefold() == ".docx":
        try:
            with zipfile.ZipFile(path) as archive:
                for name in ("docProps/core.xml", "docProps/app.xml", "docProps/custom.xml"):
                    if name in archive.namelist():
                        metadata[name] = archive.read(name).decode("utf-8", errors="replace")
        except (OSError, zipfile.BadZipFile):
            return ""
    elif path.suffix.casefold() == ".pdf":
        try:
            from pypdf import PdfReader

            raw = PdfReader(str(path)).metadata or {}
            metadata = {str(key): str(value or "") for key, value in raw.items()}
        except Exception:
            return ""
    else:
        return ""
    return sha256_text(json.dumps(metadata, ensure_ascii=False, sort_keys=True))


def discover_cv_cl(package: Path) -> dict[str, Path | None]:
    package = Path(package)
    files = [path for path in sorted(package.iterdir()) if path.is_file()]

    def choose(kind: str) -> Path | None:
        if kind == "cv":
            patterns = ("cv", "resume")
        else:
            patterns = ("cover letter", "cover_letter", "cl")
        candidates = [path for path in files if path.suffix.casefold() in {".docx", ".txt", ".md", ".pdf"} and any(token in path.name.casefold() for token in patterns)]
        # Prefer editable text over PDF for semantic auditing; PDF is fallback.
        candidates.sort(key=lambda item: (item.suffix.casefold() == ".pdf", item.name.casefold()))
        return candidates[0] if candidates else None

    return {"cv": choose("cv"), "cover_letter": choose("cover_letter")}


def material_texts(package: Path) -> dict[str, str]:
    """Read the canonical source when present; legacy DOCX/TXT is fallback.

    This makes a semantic audit independent from ZIP metadata and guarantees
    that rendering never becomes a new content-authoring step.
    """

    package = Path(package)
    try:
        from tools.workflow.materials_draft import canonical_material_texts, load_canonical_draft

        if load_canonical_draft(package):
            return canonical_material_texts(package)
    except (ImportError, OSError, ValueError, TypeError):
        pass
    output: dict[str, str] = {}
    for label, path in discover_cv_cl(package).items():
        if path is not None:
            output[label] = read_material_text(path)
    return output


def semantic_material_hashes(package: Path) -> dict[str, str]:
    return {
        label: sha256_text(normalize_text(text))
        for label, text in material_texts(package).items()
        if str(text or "").strip()
    }


def material_metadata_hashes(package: Path) -> dict[str, str]:
    result: dict[str, str] = {}
    for label, path in discover_cv_cl(package).items():
        if path is not None:
            result[label] = metadata_hash(path)
    return result


def material_container_hashes(package: Path) -> dict[str, str]:
    result: dict[str, str] = {}
    for label, path in discover_cv_cl(package).items():
        if path is not None:
            result[label] = container_hash(path)
    return result


def audit_coverage_dispositions(package: Path) -> dict[str, str]:
    """Return normalized internal JD coverage decisions visible to audit."""

    package = Path(package)
    for path in (package / "materials_draft.canonical.json", package / "materials_plan.validated.json"):
        try:
            value = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        raw = value.get("coverage_dispositions") if isinstance(value, dict) else None
        if isinstance(raw, dict):
            return {
                str(anchor_id): str(disposition)
                for anchor_id, disposition in sorted(raw.items(), key=lambda item: str(item[0]))
                if str(anchor_id).strip() and str(disposition).strip()
            }
    return {}


def audit_input_fingerprint(
    *,
    package: Path,
    jd_text: str,
    rules_digest: str,
    lessons_digest: str = "",
    claim_contract: dict[str, Any] | None = None,
) -> str:
    """Fingerprint the inputs visible to the presentation auditor.

    ``claim_contract`` remains an ignored compatibility argument for callers
    from the pre-v2 API.  It is intentionally not represented in the digest:
    a contract/authorization change must not invalidate or influence the new
    JD-mapping and display-quality audit.
    """

    del claim_contract
    payload = {
        "jd_sha256": sha256_text(normalize_text(jd_text)),
        "semantic_material_hashes": semantic_material_hashes(package),
        "rules_digest": rules_digest,
        "lessons_digest": lessons_digest,
        "coverage_dispositions": audit_coverage_dispositions(package),
    }
    return sha256_text(json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")))


def audit_context_fingerprint(
    *,
    jd_text: str,
    rules_digest: str,
    lessons_digest: str = "",
    claim_contract: dict[str, Any] | None = None,
    coverage_dispositions: dict[str, str] | None = None,
) -> str:
    """Fingerprint stable inputs while deliberately excluding draft text.

    A finding-scoped repair changes the draft but remains in the same bounded
    audit run.  Excluding draft text here prevents each repair from resetting
    the three-attempt budget.
    """

    del claim_contract
    payload = {
        "jd_sha256": sha256_text(normalize_text(jd_text)),
        "rules_digest": rules_digest,
        "lessons_digest": lessons_digest,
        "coverage_dispositions": dict(coverage_dispositions or {}),
    }
    return sha256_text(json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")))


def metadata_only_change(previous: dict[str, Any], package: Path) -> bool:
    return bool(previous.get("semantic_material_hashes")) and previous.get("semantic_material_hashes") == semantic_material_hashes(package) and previous.get("metadata_hashes") != material_metadata_hashes(package)
