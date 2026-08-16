"""Render audited canonical CV/CL content into DOCX and PDF artifacts."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
from datetime import datetime, timezone
from math import ceil
from pathlib import Path
from typing import Any
from uuid import uuid4

from tools.io_utils import atomic_write_json
from tools.job_materials.manifest import load_job_manifest
from tools.job_materials.paths import find_latest_cl_master_docx, find_latest_master_docx
from tools.workflow.materials_draft import canonical_digest, canonical_material_texts, load_canonical_draft
from tools.workflow.materials_hashes import container_hash, normalize_text, read_material_text, semantic_material_hashes
from tools.workflow.materials_metadata import metadata_violations, sanitize_docx_metadata

RENDER_RECEIPT_NAME = "materials_render_receipt.json"
FORMAT_REPORT_NAME = "materials_format_report.json"
RENDERER_VERSION = "canonical-template-docx-v3"

TEMPLATE_STYLES: dict[str, tuple[str, ...]] = {
    "cv": ("Resume Section", "Job Heading", "Resume Bullet", "Compact Line"),
    "cover_letter": ("Letter Body", "Letter Bullet", "Letter Compact"),
}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _load(path: Path) -> dict[str, Any]:
    try:
        value = json.loads(Path(path).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _candidate_name(workspace: Path) -> str:
    config = _load(Path(workspace) / "00_Profile" / "config.personal.json")
    value = str(config.get("candidate_name") or config.get("name") or "").strip()
    if not value:
        raise ValueError("candidate_name_missing")
    return value


def _filename_component(value: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|_]+", " ", str(value or ""))
    # JobsFlow material names use spaces and retain substantive parentheses.
    text = re.sub(r"[\-–—]+", " ", text)
    return re.sub(r"\s+", " ", text).strip(" .")


def expected_filenames(package: Path, workspace: Path) -> dict[str, str]:
    # A vNext generation freezes one entity contract. Renderer/filename code
    # must not re-parse the manifest and accidentally turn a recruiter into
    # the hiring company or revive an old long/alternate title.
    vnext_bundle = Path(package) / "materials_vnext" / "current_job_bundle.json"
    if vnext_bundle.is_file():
        try:
            value = json.loads(vnext_bundle.read_text(encoding="utf-8"))
            entity = value.get("entity") if isinstance(value.get("entity"), dict) else {}
            role_value = str(entity.get("role_primary") or "").strip()
            target_value = str(entity.get("employer_name") or "").strip()
            if not target_value and str(entity.get("recruiter_boundary") or "") == "direct_employer":
                target_value = str(entity.get("application_target") or "").strip()
            if str(entity.get("recruiter_boundary") or "").startswith("recruiter") and not target_value:
                target_value = ""
            if role_value:
                role = _filename_component(role_value)
                company = _filename_component(target_value)
                candidate = _filename_component(_candidate_name(Path(workspace)))
                prefix = " ".join(item for item in (candidate, company, role) if item).strip()
                return {
                    "cv_docx": f"{prefix} CV.docx",
                    "cl_docx": f"{prefix} Cover Letter.docx",
                    "cv_pdf": f"{prefix} CV.pdf",
                    "cl_pdf": f"{prefix} Cover Letter.pdf",
                }
        except (OSError, json.JSONDecodeError, TypeError, ValueError):
            pass
    manifest = load_job_manifest(Path(package)) or {}
    job = manifest.get("job") if isinstance(manifest.get("job"), dict) else {}
    role = _filename_component(str(job.get("role_material") or job.get("role_display") or "Application"))
    company = _filename_component(str(job.get("company_out") or job.get("employer_name") or ""))
    candidate = _filename_component(_candidate_name(Path(workspace)))
    prefix = " ".join(item for item in (candidate, company, role) if item).strip()
    return {
        "cv_docx": f"{prefix} CV.docx",
        "cl_docx": f"{prefix} Cover Letter.docx",
        "cv_pdf": f"{prefix} CV.pdf",
        "cl_pdf": f"{prefix} Cover Letter.pdf",
    }


def _audit_current(package: Path) -> bool:
    report = _load(Path(package) / "materials_audit.json")
    return bool(
        report.get("status") == "passed"
        and report.get("content_gate") == "passed"
        and isinstance(report.get("semantic_material_hashes"), dict)
        and report.get("semantic_material_hashes") == semantic_material_hashes(Path(package))
        and int((report.get("open_counts") or {}).get("P0", 0)) == 0
        and int((report.get("open_counts") or {}).get("P1", 0)) == 0
    )


def _archive_previous_outputs(package: Path, names: dict[str, str]) -> None:
    active = {value for value in names.values()}
    stale = [
        path
        for path in Path(package).iterdir()
        if path.is_file()
        and path.suffix.casefold() in {".docx", ".pdf", ".txt"}
        and (
            "cv" in path.name.casefold()
            or "cover" in path.name.casefold()
            or path.name.casefold() in {"cl.txt", "cover_letter.txt"}
        )
        and path.name not in active
    ]
    if not stale:
        return
    history = Path(package) / ".history" / f"render-{uuid4().hex[:10]}"
    history.mkdir(parents=True, exist_ok=True)
    for path in stale:
        shutil.move(str(path), str(history / path.name))
        sidecar = path.with_suffix(path.suffix + ".jobsflow.json")
        if sidecar.is_file():
            shutil.move(str(sidecar), str(history / sidecar.name))


def _set_page(document) -> None:
    from docx.enum.section import WD_SECTION
    from docx.shared import Cm

    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.25)
        section.bottom_margin = Cm(1.20)
        section.left_margin = Cm(1.45)
        section.right_margin = Cm(1.45)


def _template_document(path: Path, *, material: str):
    from docx import Document

    path = Path(path)
    if not path.is_file():
        raise ValueError(f"base_template_missing:{material}")
    document = Document(str(path))
    available = {str(style.name) for style in document.styles}
    missing = [name for name in TEMPLATE_STYLES[material] if name not in available]
    if missing:
        raise ValueError(f"base_template_invalid:{material}:missing_styles={','.join(missing)}")
    return document


def _clear_document_body(document) -> None:
    """Remove master content while retaining sections, styles and page setup."""

    body = document._element.body
    for child in list(body):
        # sectPr stores the master's page size, margins and header/footer
        # relationship. It must survive the content replacement.
        if child.tag.rsplit("}", 1)[-1] == "sectPr":
            continue
        body.remove(child)


def _paragraph_prototype(paragraph) -> dict[str, Any]:
    """Capture the *direct* formatting of one master paragraph.

    A paragraph style is not enough to reproduce a Word master.  The lane
    masters intentionally keep the visual contract (colour, weight, tabs,
    numbering and spacing) as direct OOXML on representative paragraphs.  We
    therefore copy those XML fragments rather than trying to reconstruct the
    appearance from style names.
    """

    return {
        "style": str(paragraph.style.name),
        "ppr": deepcopy(paragraph._p.pPr) if paragraph._p.pPr is not None else None,
        "rprs": [deepcopy(run._r.rPr) if run._r.rPr is not None else None for run in paragraph.runs],
        "text": str(paragraph.text or ""),
    }


def _first_paragraph(document, predicate) -> dict[str, Any] | None:
    for paragraph in document.paragraphs:
        if predicate(paragraph):
            return _paragraph_prototype(paragraph)
    return None


def _template_prototypes(document, *, material: str) -> dict[str, dict[str, Any]]:
    """Return named visual prototypes from the selected lane master."""

    paragraphs = list(document.paragraphs)
    normal = [p for p in paragraphs if str(p.style.name) == "Normal"]
    prototypes: dict[str, dict[str, Any]] = {}

    # Both the CV and CL masters use the first two Normal paragraphs for the
    # candidate name and contact line.  Do not depend on the canonical block
    # type: a contact block is the title in the actual runtime draft.
    if normal:
        prototypes["title"] = _paragraph_prototype(normal[0])
    if len(normal) > 1:
        prototypes["contact"] = _paragraph_prototype(normal[1])

    if material == "cv":
        for key, style in (
            ("section", "Resume Section"),
            ("job_heading", "Job Heading"),
            ("bullet", "Resume Bullet"),
            ("core", "Compact Line"),
        ):
            match = next((p for p in paragraphs if str(p.style.name) == style), None)
            if match is not None:
                prototypes[key] = _paragraph_prototype(match)
        compact = next(
            (
                p
                for p in paragraphs
                if str(p.style.name) == "Compact Line"
                and len(p.runs) >= 2
                and not any(str(run.font.color.rgb or "").casefold() in {"17365d", "17365d00"} for run in p.runs)
            ),
            None,
        )
        if compact is not None:
            prototypes["compact"] = _paragraph_prototype(compact)
        else:
            # A minimal test master may only expose the core prototype.  Real
            # lane masters contain a separate black Compact Line prototype;
            # using core as a safe fallback keeps the template contract
            # backwards-compatible while still preserving its direct XML.
            prototypes["compact"] = prototypes["core"]
        summary = next(
            (p for p in paragraphs if str(p.style.name) == "Normal" and p not in normal[:2]),
            None,
        )
        if summary is not None:
            prototypes["summary"] = _paragraph_prototype(summary)
    else:
        for key, style in (
            ("body", "Letter Body"),
            ("bullet", "Letter Bullet"),
            ("compact", "Letter Compact"),
        ):
            match = next((p for p in paragraphs if str(p.style.name) == style), None)
            if match is not None:
                prototypes[key] = _paragraph_prototype(match)
        subject = next((p for p in normal if str(p.text).lstrip().startswith("Re:")), None)
        if subject is not None:
            prototypes["subject"] = _paragraph_prototype(subject)

    missing = [key for key in ("title", "contact") if key not in prototypes]
    if material == "cv":
        missing.extend(key for key in ("section", "job_heading", "bullet", "core", "summary", "compact") if key not in prototypes)
    else:
        missing.extend(key for key in ("body", "bullet", "compact", "subject") if key not in prototypes)
    if missing:
        raise ValueError(f"base_template_invalid:{material}:missing_prototypes={','.join(missing)}")
    return prototypes


def _new_paragraph(document, prototype: dict[str, Any]):
    """Create a paragraph and copy prototype pPr (tabs/numbering/spacing)."""

    paragraph = document.add_paragraph(style=str(prototype["style"]))
    ppr = paragraph._p.get_or_add_pPr()
    # Keep the style binding created by python-docx, but replace every other
    # paragraph property with the master property.  In particular this keeps
    # Job Heading's right tab and Resume Bullet/Letter Bullet's numPr.
    style_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"
    for child in list(ppr):
        if child.tag != style_tag:
            ppr.remove(child)
    source = prototype.get("ppr")
    if source is not None:
        for child in source:
            if child.tag != style_tag:
                ppr.append(deepcopy(child))
    return paragraph


def _add_run(paragraph, text: str, prototype: dict[str, Any] | None, run_index: int = 0):
    run = paragraph.add_run(str(text or ""))
    if prototype is None:
        return run
    rprs = prototype.get("rprs") or []
    source = rprs[min(run_index, len(rprs) - 1)] if rprs else None
    if source is None:
        return run
    rpr = run._r.get_or_add_rPr()
    for child in list(rpr):
        rpr.remove(child)
    for child in source:
        rpr.append(deepcopy(child))
    return run


def _split_label(text: str) -> tuple[str, str] | None:
    # Core-expertise and CL pillar prototypes use a coloured/bold label run
    # followed by a normal evidence run.  Preserve the delimiter in the label.
    match = re.match(r"^(.+?)\s+-\s+(.+)$", text.strip())
    if match:
        return match.group(1).rstrip(), match.group(2).lstrip()
    match = re.match(r"^(.+?:)\s*(.+)$", text.strip())
    if match:
        return match.group(1), match.group(2)
    return None


def _job_heading_parts(text: str) -> tuple[str, str | None]:
    value = text.strip()
    # Baseline compilation normalises ordinary whitespace, so a heading may
    # arrive with one space instead of the master's tab/double-space
    # delimiter.  Recognise a terminal date range independently of that
    # delimiter and keep the renderer responsible for restoring the master's
    # right-aligned tab.  This is intentionally bounded to common CV date
    # forms; it must not mistake a year in a role title for a date column.
    date = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}"
    year = r"\d{4}"
    date_suffix = re.compile(
        rf"\s+(?P<date>(?:{date}|{year})\s*(?:-|–|—|to)\s*(?:Present|Current|{date}|{year})|(?:{date}|{year}))$",
        re.IGNORECASE,
    )
    match = date_suffix.search(value)
    if match and match.start() > 0:
        return value[: match.start()].rstrip(), match.group("date").strip()
    parts = re.split(r"\s{2,}|\t", value, maxsplit=1)
    if len(parts) == 2 and parts[1].strip():
        return parts[0].strip(), parts[1].strip()
    return value, None


def _layout_units(document, *, material: str) -> float:
    """Estimate wrapped-line demand without opening Word or a PDF engine.

    This is deliberately a routing heuristic, not a page-count assertion.  It
    uses the same template-bound styles as the renderer and lets us detect the
    uncommon case where a much shorter customised document leaves an
    unbalanced lower half of the page.
    """

    widths = {
        "cv": {
            "Normal": 92,
            "Resume Section": 96,
            "Job Heading": 84,
            "Resume Bullet": 94,
            "Compact Line": 104,
        },
        "cover_letter": {
            "Normal": 88,
            "Letter Body": 88,
            "Letter Bullet": 88,
            "Letter Compact": 100,
        },
    }[material]
    total = 0.0
    for paragraph in document.paragraphs:
        text = str(paragraph.text or "")
        if not text:
            total += 0.25
            continue
        width = widths.get(str(paragraph.style.name), 90)
        total += sum(max(1, ceil(len(part) / width)) for part in text.split("\n"))
    return total


def _apply_visual_balance(
    document,
    *,
    material: str,
    target_units: float,
    target_paragraphs: int,
) -> dict[str, Any]:
    """Add bounded inter-block rhythm when a short draft underfills a master.

    The lane master remains authoritative for fonts, colours, margins, tabs,
    numbering and line spacing.  This helper only adds modest space *between*
    major content blocks when the canonical draft is materially shorter than
    the lane baseline.  It never stretches text, invents content, or changes a
    one-page decision; the normal PDF/page gate remains the final authority.
    """

    actual_units = _layout_units(document, material=material)
    gap_units = max(0.0, target_units - actual_units)
    paragraph_gap = max(0, int(target_paragraphs) - len(document.paragraphs))
    if gap_units < 1.0 and paragraph_gap < 2:
        return {
            "mode": "template_native",
            "target_units": round(target_units, 2),
            "actual_units": round(actual_units, 2),
            "gap_units": round(gap_units, 2),
            "extra_space_after_pt": 0.0,
        }

    if material == "cv":
        candidates = [
            paragraph
            for paragraph in document.paragraphs
            if str(paragraph.style.name) in {"Resume Section", "Job Heading"}
        ]
    else:
        candidates = [
            paragraph
            for paragraph in document.paragraphs
            if str(paragraph.style.name) in {"Letter Compact", "Letter Body", "Letter Bullet"}
            and str(paragraph.text or "").strip()
        ][2:]  # leave the name/contact header untouched
    if not candidates:
        return {
            "mode": "template_native",
            "target_units": round(target_units, 2),
            "actual_units": round(actual_units, 2),
            "gap_units": round(gap_units, 2),
            "extra_space_after_pt": 0.0,
        }

    # Paragraph-count differences capture omitted optional template slots (for
    # example a shorter recipient block); wrapped-line differences capture
    # fewer experience bullets.  Keep the adjustment bounded so a sparse draft
    # never becomes artificially airy.
    target_points = min(90.0, max(gap_units * 10.5, paragraph_gap * 14.0))
    per_paragraph = min(7.0, target_points / len(candidates))
    for paragraph in candidates:
        current = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0.0
        from docx.shared import Pt

        paragraph.paragraph_format.space_after = Pt(current + per_paragraph)
    return {
        "mode": "adaptive_inter_block_rhythm",
        "target_units": round(target_units, 2),
        "actual_units": round(actual_units, 2),
        "gap_units": round(gap_units, 2),
        "paragraph_gap": paragraph_gap,
        "adjusted_blocks": len(candidates),
        "extra_space_after_pt": round(per_paragraph, 2),
    }


def _add_block(
    document,
    block: dict[str, Any],
    *,
    material: str,
    position: int,
    prototypes: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    block_type = str(block.get("type") or "paragraph")
    text = str(block.get("text") or "").strip()
    section = str(block.get("section") or "").casefold()
    presentation_role = str(block.get("presentation_role") or "").casefold()
    source_style = str(block.get("source_style") or "")

    if material == "cv":
        if position == 0 and block_type in {"contact", "heading"}:
            paragraph = _new_paragraph(document, prototypes["title"])
            _add_run(paragraph, text, prototypes["title"])
        elif block_type == "contact":
            paragraph = _new_paragraph(document, prototypes["contact"])
            _add_run(paragraph, text, prototypes["contact"])
        elif block_type == "heading" and section in {"summary", "core", "experience", "education", "qualifications"} and not block.get("experience_id"):
            paragraph = _new_paragraph(document, prototypes["section"])
            _add_run(paragraph, text, prototypes["section"])
        elif block_type == "heading" and section == "experience":
            paragraph = _new_paragraph(document, prototypes["job_heading"])
            role, date = _job_heading_parts(text)
            _add_run(paragraph, role, prototypes["job_heading"], 0)
            if date:
                _add_run(paragraph, "\t" + date, prototypes["job_heading"], 1)
        elif (presentation_role == "core_line" or (source_style == "Compact Line" and section == "core")):
            paragraph = _new_paragraph(document, prototypes["core"])
            label_body = _split_label(text)
            if label_body:
                _add_run(paragraph, label_body[0], prototypes["core"], 0)
                separator = " - " if re.search(r"\s+-\s+", text) else " "
                _add_run(paragraph, separator + label_body[1], prototypes["core"], 1)
            else:
                _add_run(paragraph, text, prototypes["core"], 1)
        # Education and Qualifications are compact master lines, not
        # experience bullets.  This branch deliberately precedes the generic
        # bullet branch so a baseline block cannot lose its section styling
        # merely because its semantic type is ``bullet`` for compatibility.
        elif presentation_role == "compact_line" or source_style == "Compact Line" or section in {"education", "qualifications", "compact", "contact"}:
            paragraph = _new_paragraph(document, prototypes["compact"])
            if section == "education" and ", " in text:
                label, body = text.split(", ", 1)
                _add_run(paragraph, label, prototypes["compact"], 0)
                _add_run(paragraph, ", " + body, prototypes["compact"], 1)
            elif section == "qualifications" and ": " in text:
                label, body = text.split(": ", 1)
                _add_run(paragraph, label + ": ", prototypes["compact"], 0)
                _add_run(paragraph, body, prototypes["compact"], 1)
            else:
                _add_run(paragraph, text, prototypes["compact"], 1)
        elif block_type == "bullet":
            paragraph = _new_paragraph(document, prototypes["bullet"])
            _add_run(paragraph, text, prototypes["bullet"])
        elif section == "summary":
            paragraph = _new_paragraph(document, prototypes["summary"])
            _add_run(paragraph, text, prototypes["summary"])
        else:
            paragraph = _new_paragraph(document, prototypes["summary"])
            _add_run(paragraph, text, prototypes["summary"])
    else:
        if block_type == "contact" and position == 0:
            paragraph = _new_paragraph(document, prototypes["title"])
            _add_run(paragraph, text, prototypes["title"])
        elif block_type == "contact":
            paragraph = _new_paragraph(document, prototypes["contact"])
            _add_run(paragraph, text, prototypes["contact"])
        elif section == "subject" or text.casefold().startswith("re: application"):
            paragraph = _new_paragraph(document, prototypes["subject"])
            _add_run(paragraph, text, prototypes["subject"])
        elif block_type == "bullet" or section == "pillar":
            paragraph = _new_paragraph(document, prototypes["bullet"])
            label_body = _split_label(text)
            if label_body:
                _add_run(paragraph, label_body[0], prototypes["bullet"], 0)
                _add_run(paragraph, " " + label_body[1], prototypes["bullet"], 1)
            else:
                _add_run(paragraph, text, prototypes["bullet"])
        elif block_type == "signoff" or section in {"date", "recipient", "signoff"}:
            paragraph = _new_paragraph(document, prototypes["compact"])
            _add_run(paragraph, text, prototypes["compact"])
        else:
            paragraph = _new_paragraph(document, prototypes["body"])
            _add_run(paragraph, text, prototypes["body"])

    # Do not impose a second, model-dependent layout system.  The paragraph
    # prototype copied above is the sole formatting authority; these flags are
    # harmless accessibility defaults and do not overwrite colour, tabs or
    # numbering from the master.
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True


def _render_document(
    blocks: list[dict[str, Any]],
    path: Path,
    *,
    material: str,
    title: str,
    author: str,
    template: Path,
) -> None:
    document = _template_document(template, material=material)
    prototypes = _template_prototypes(document, material=material)
    target_units = _layout_units(document, material=material)
    target_paragraphs = len(document.paragraphs)
    _clear_document_body(document)
    for position, block in enumerate(blocks):
        _add_block(document, block, material=material, position=position, prototypes=prototypes)
    layout_balance = _apply_visual_balance(
        document,
        material=material,
        target_units=target_units,
        target_paragraphs=target_paragraphs,
    )
    document.save(str(path))
    sanitize_docx_metadata(path, title=title, subject="Job application CV" if material == "cv" else "Job application Cover Letter", author=author)
    return layout_balance


def _template_paths(package: Path, workspace: Path) -> dict[str, Path]:
    manifest = load_job_manifest(Path(package)) or {}
    lane = ""
    bundle_path = Path(package) / "materials_vnext" / "current_job_bundle.json"
    if bundle_path.is_file():
        try:
            bundle = json.loads(bundle_path.read_text(encoding="utf-8"))
            lane = str(bundle.get("lane") or "").upper()
        except (OSError, json.JSONDecodeError, TypeError):
            lane = ""
    lane = lane or str(manifest.get("lane") or Path(package).parts[-3][:1] or "").upper()
    cv = find_latest_master_docx(lane, workspace)
    cl = find_latest_cl_master_docx(lane, workspace)
    if cv is None:
        raise ValueError("base_template_missing:cv")
    if cl is None:
        raise ValueError("base_template_missing:cover_letter")
    # Validate styles before any existing output is archived. This ensures a
    # bad or generic master cannot destroy a previously valid package.
    _template_document(cv, material="cv")
    _template_document(cl, material="cover_letter")
    return {"cv": cv.resolve(), "cover_letter": cl.resolve()}


def _receipt_current(
    package: Path,
    names: dict[str, str],
    canonical_sha256: str,
    template_paths: dict[str, Path],
) -> bool:
    receipt = _load(Path(package) / RENDER_RECEIPT_NAME)
    if receipt.get("canonical_sha256") != canonical_sha256 or receipt.get("renderer_version") != RENDERER_VERSION:
        return False
    expected_templates = {key: container_hash(path) for key, path in template_paths.items()}
    if receipt.get("template_sha256") != expected_templates:
        return False
    for key in ("cv_docx", "cl_docx"):
        path = Path(package) / names[key]
        if not path.is_file() or (receipt.get("docx_hashes") or {}).get(key) != container_hash(path):
            return False
    return True


def render_canonical_docx(package: Path, workspace: Path, *, force: bool = False) -> dict[str, Any]:
    package = Path(package)
    draft = load_canonical_draft(package)
    if not draft:
        raise ValueError("canonical_draft_missing")
    if (
        draft.get("baseline_sha256")
        and str(draft.get("artifact_type") or "") != "jobsflow_canonical_cv_cl"
        and str(draft.get("compiled_from") or "") != "bounded_baseline_transform"
    ):
        raise ValueError("baseline_transform_required")
    if draft.get("baseline_sha256") and str(draft.get("artifact_type") or "") != "jobsflow_canonical_cv_cl":
        from tools.workflow.materials_baseline import load_content_baseline, validate_content_floor

        floor_errors = validate_content_floor(load_content_baseline(package), draft)
        if floor_errors:
            raise ValueError("baseline_content_floor_invalid:" + ",".join(floor_errors))
    if not _audit_current(package):
        raise ValueError("content_audit_not_current")
    names = expected_filenames(package, workspace)
    templates = _template_paths(package, workspace)
    digest = str(draft.get("canonical_sha256") or canonical_digest(draft))
    if not force and _receipt_current(package, names, digest, templates):
        return {"status": "cached", "filenames": names, "receipt": _load(package / RENDER_RECEIPT_NAME)}
    _archive_previous_outputs(package, names)
    # The frozen outbound manifest belongs to the previous derived-artifact
    # generation.  Canonical text + independent audit are the source of truth;
    # the final validator freezes the newly rendered generation after PDF and
    # format gates pass.  Keeping legacy hashes here creates a permanent
    # stale_artifact loop.
    old_manifest = package / "artifact_hashes.json"
    if old_manifest.is_file():
        old_manifest.unlink()
    candidate = _candidate_name(Path(workspace))
    cv = package / names["cv_docx"]
    cl = package / names["cl_docx"]
    cv_layout = _render_document(
        list((draft.get("cv") or {}).get("blocks") or []),
        cv,
        material="cv",
        title=cv.stem,
        author=candidate,
        template=templates["cv"],
    )
    cl_layout = _render_document(
        list((draft.get("cover_letter") or {}).get("blocks") or []),
        cl,
        material="cover_letter",
        title=cl.stem,
        author=candidate,
        template=templates["cover_letter"],
    )
    receipt = {
        "schema_version": 1,
        "renderer_version": RENDERER_VERSION,
        "canonical_sha256": digest,
        "filenames": names,
        "docx_hashes": {"cv_docx": container_hash(cv), "cl_docx": container_hash(cl)},
        "template_paths": {key: str(path) for key, path in templates.items()},
        "template_sha256": {key: container_hash(path) for key, path in templates.items()},
        "template_style_contract": TEMPLATE_STYLES,
        "layout_balance": {"cv": cv_layout, "cover_letter": cl_layout},
        "rendered_at": _now(),
    }
    atomic_write_json(package / RENDER_RECEIPT_NAME, receipt)
    return {"status": "rendered", "filenames": names, "receipt": receipt}


def convert_rendered_pdfs(
    package: Path,
    workspace: Path,
    *,
    engine: str = "libreoffice",
    force: bool = False,
    parallel: bool = True,
) -> dict[str, Any]:
    from tools.fresh_24h.docx_to_pdf import convert

    package = Path(package)
    render = render_canonical_docx(package, workspace, force=False)
    names = expected_filenames(package, workspace)
    pairs = [
        (package / names["cv_docx"], package / names["cv_pdf"]),
        (package / names["cl_docx"], package / names["cl_pdf"]),
    ]

    def run_pair(pair):
        docx, pdf = pair
        # render_canonical_docx already normalizes core properties and records
        # the resulting DOCX hash in materials_render_receipt.json.  Running
        # sanitize_docx_metadata again here rewrites the ZIP and invalidates
        # that receipt before the format gate can read it.  Conversion is
        # therefore a pure DOCX -> PDF step; any metadata change must go
        # through render and produce a new receipt.
        return convert(docx, pdf, engine=engine, force=force, sanitize_metadata=False)

    if parallel:
        with ThreadPoolExecutor(max_workers=2, thread_name_prefix="jobsflow-pdf") as pool:
            pdfs = list(pool.map(run_pair, pairs))
    else:
        pdfs = [run_pair(pair) for pair in pairs]
    return {"status": "converted", "render": render, "pdfs": [str(path) for path in pdfs], "filenames": names}


def _pdf_stats(path: Path) -> tuple[int, str]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return len(reader.pages), "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return 0, ""


def _render_matches_canonical(package: Path, names: dict[str, str]) -> list[str]:
    expected = canonical_material_texts(package)
    errors: list[str] = []
    for material, key in (("cv", "cv_docx"), ("cover_letter", "cl_docx")):
        live = normalize_text(read_material_text(Path(package) / names[key])).casefold()
        for line in normalize_text(expected.get(material, "")).splitlines():
            if line.casefold() not in live:
                errors.append(f"rendered_text_mismatch:{material}")
                break
    return errors


def _template_style_findings(package: Path, names: dict[str, str]) -> list[dict[str, str]]:
    """Reject a DOCX that was produced outside the fixed lane template chain."""

    receipt = _load(Path(package) / RENDER_RECEIPT_NAME)
    findings: list[dict[str, str]] = []
    if receipt.get("renderer_version") != RENDERER_VERSION:
        findings.append({"code": "template_binding_missing", "artifact": "docx", "evidence": "render receipt is not template-bound"})
    template_paths = receipt.get("template_paths") if isinstance(receipt.get("template_paths"), dict) else {}
    recorded_templates = receipt.get("template_sha256") if isinstance(receipt.get("template_sha256"), dict) else {}
    if template_paths and not recorded_templates:
        findings.append({
            "code": "template_receipt_hash_missing",
            "artifact": "docx",
            "evidence": "render receipt has no lane-master hash binding",
        })
    for material, raw_path in template_paths.items():
        template_path = Path(str(raw_path))
        if not template_path.is_file():
            findings.append({"code": "template_binding_missing", "artifact": material, "evidence": "template source unavailable"})
            continue
        recorded = str(recorded_templates.get(material) or "")
        live = container_hash(template_path)
        if recorded and recorded != live:
            findings.append({
                "code": "template_changed_after_render",
                "artifact": material,
                "evidence": "lane master changed after DOCX render; rerender through the fixed gateway",
            })
    for material, key in (("cv", "cv_docx"), ("cover_letter", "cl_docx")):
        path = Path(package) / names[key]
        if not path.is_file():
            continue
        try:
            from docx import Document

            document = Document(str(path))
            available = {str(style.name) for style in document.styles}
        except Exception:
            findings.append({"code": "template_style_unreadable", "artifact": key, "evidence": path.name})
            continue
        missing = [name for name in TEMPLATE_STYLES[material] if name not in available]
        if missing:
            findings.append(
                {
                    "code": "template_style_contract_failed",
                    "artifact": key,
                    "evidence": ",".join(missing),
                }
            )
            continue

        template_path = Path(str(template_paths.get(material) or ""))
        if not template_path.is_file():
            findings.append({"code": "template_binding_missing", "artifact": key, "evidence": "template path unavailable"})
            continue
        try:
            template = Document(str(template_path))
            prototypes = _template_prototypes(template, material=material)
        except Exception as exc:
            findings.append({"code": "template_prototype_unreadable", "artifact": key, "evidence": str(exc)})
            continue

        def xml_signature(element) -> tuple[Any, ...]:
            """Compare OOXML properties without namespace-prefix noise."""

            if element is None:
                return ()

            def walk(node):
                local = str(node.tag).rsplit("}", 1)[-1]
                attrs = tuple(sorted((str(key).rsplit("}", 1)[-1], str(value)) for key, value in node.attrib.items()))
                children = tuple(walk(child) for child in node)
                return (local, attrs, str(node.text or ""), children)

            return walk(element)

        def rpr_signature(run) -> tuple[Any, ...]:
            return xml_signature(run._r.rPr if run is not None and run._r.rPr is not None else None)

        def prototype_rpr_signature(name: str, index: int = 0) -> str:
            runs = prototypes[name].get("rprs") or []
            source = runs[min(index, len(runs) - 1)] if runs else None
            return xml_signature(source)

        def has_ppr_tag(paragraph, local_name: str) -> bool:
            ppr = paragraph._p.pPr
            return ppr is not None and ppr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + local_name) is not None

        def prototype_has_ppr_tag(name: str, local_name: str) -> bool:
            ppr = prototypes[name].get("ppr")
            return ppr is not None and ppr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + local_name) is not None

        def direct_ppr_signature(element) -> tuple[Any, ...]:
            """Compare layout-bearing direct properties, excluding spacing.

            Visual-balance compensation intentionally adjusts ``w:spacing``
            between blocks.  Tabs, numbering, indentation and alignment are
            template-owned and must remain byte-equivalent to the prototype.
            """

            if element is None:
                return ()
            wanted = {"tabs", "numPr", "ind", "jc", "keepNext", "pageBreakBefore"}
            return tuple(xml_signature(child) for child in element if str(child.tag).rsplit("}", 1)[-1] in wanted)

        def prototype_direct_ppr_signature(name: str) -> tuple[Any, ...]:
            return direct_ppr_signature(prototypes[name].get("ppr"))

        def expected_run_index(prototype_name: str, actual_index: int, actual_count: int) -> int:
            runs = prototypes[prototype_name].get("rprs") or []
            if not runs:
                return 0
            # Compact/core one-run blocks are emitted with the evidence/body
            # run (index 1) by _add_block; split label/body blocks preserve
            # their corresponding prototype run indices.
            if actual_count == 1 and prototype_name in {"compact", "core"} and len(runs) > 1:
                return 1
            return min(actual_index, len(runs) - 1)

        def expected_prototype(block: dict[str, Any], position: int) -> str:
            """Map a canonical block to the host-owned master prototype.

            This mapping is deliberately derived from baseline metadata first
            and only falls back to the legacy semantic fields for old fixture
            drafts.  A model cannot select a style by changing its prose.
            """

            block_type = str(block.get("type") or "paragraph")
            section_name = str(block.get("section") or "").casefold()
            presentation = str(block.get("presentation_role") or "").casefold()
            source_style = str(block.get("source_style") or "")
            if position == 0 and block_type in {"contact", "heading"}:
                return "title"
            if block_type == "contact":
                return "contact"
            if material == "cv":
                if presentation == "target_role" or section_name == "target_role":
                    return "summary"
                if presentation == "section_heading" or (
                    block_type == "heading"
                    and section_name in {"summary", "core", "experience", "education", "qualifications"}
                    and not block.get("experience_id")
                ):
                    return "section"
                if presentation == "job_heading" or (block_type == "heading" and section_name == "experience"):
                    return "job_heading"
                if presentation == "core_line" or (source_style == "Compact Line" and section_name == "core"):
                    return "core"
                if presentation == "compact_line" or source_style == "Compact Line" or section_name in {"education", "qualifications", "compact", "contact"}:
                    return "compact"
                if presentation == "experience_bullet" or block_type == "bullet":
                    return "bullet"
                if section_name == "summary":
                    return "summary"
                return "summary"
            if presentation == "subject" or section_name == "subject" or str(block.get("text") or "").lstrip().casefold().startswith("re:"):
                return "subject"
            if block_type == "contact":
                return "contact"
            if block_type == "bullet" or section_name == "pillar":
                return "bullet"
            if block_type == "signoff" or section_name in {"date", "recipient", "signoff"}:
                return "compact"
            return "body"

        # The style name check below is not enough: a generic model can copy
        # the style catalogue while putting every paragraph in Normal.  Walk
        # every canonical block in order and compare the live paragraph to the
        # exact lane prototype.  This catches section/education/date drift in
        # all jobs, not just the first representative paragraph.
        canonical = load_canonical_draft(Path(package))
        canonical_blocks = [
            item for item in ((canonical.get(material) or {}).get("blocks") or [])
            if isinstance(item, dict)
        ]
        if canonical_blocks and len(document.paragraphs) < len(canonical_blocks):
            findings.append({"code": "template_block_count_failed", "artifact": key, "evidence": f"expected={len(canonical_blocks)} actual={len(document.paragraphs)}"})
        for position, block in enumerate(canonical_blocks):
            if position >= len(document.paragraphs):
                break
            paragraph = document.paragraphs[position]
            prototype_name = expected_prototype(block, position)
            prototype = prototypes.get(prototype_name)
            if prototype is None:
                findings.append({"code": "template_prototype_missing", "artifact": key, "evidence": prototype_name})
                continue
            expected_style = str(prototype.get("style") or "")
            if str(paragraph.style.name) != expected_style:
                findings.append({
                    "code": "template_block_style_failed",
                    "artifact": key,
                    "evidence": f"{block.get('id')}: expected={expected_style} actual={paragraph.style.name}",
                })
            if direct_ppr_signature(paragraph._p.pPr) != prototype_direct_ppr_signature(prototype_name):
                findings.append({
                    "code": "template_block_layout_failed",
                    "artifact": key,
                    "evidence": f"{block.get('id')}: direct paragraph layout differs from lane master",
                })
            if paragraph.runs and prototype.get("rprs"):
                run_format_errors = []
                for run_index, run in enumerate(paragraph.runs):
                    expected_index = expected_run_index(prototype_name, run_index, len(paragraph.runs))
                    if rpr_signature(run) != prototype_rpr_signature(prototype_name, expected_index):
                        run_format_errors.append(str(run_index))
                if run_format_errors:
                    findings.append({
                        "code": "template_block_direct_format_failed",
                        "artifact": key,
                        "evidence": f"{block.get('id')}: prototype={prototype_name} runs={','.join(run_format_errors)}",
                    })
            if prototype_name == "job_heading":
                _, date_text = _job_heading_parts(str(block.get("text") or ""))
                if date_text:
                    if "\t" not in paragraph.text or not has_ppr_tag(paragraph, "tabs"):
                        findings.append({
                            "code": "template_job_heading_date_tab_failed",
                            "artifact": key,
                            "evidence": f"{block.get('id')}: date column is not tab-bound",
                        })
                    if len(paragraph.runs) >= 2 and rpr_signature(paragraph.runs[1]) != prototype_rpr_signature(prototype_name, 1):
                        findings.append({
                            "code": "template_job_heading_date_format_failed",
                            "artifact": key,
                            "evidence": f"{block.get('id')}: date run differs from lane master",
                        })
            if prototype_name in {"bullet", "core"} and prototype_has_ppr_tag(prototype_name, "numPr") and not has_ppr_tag(paragraph, "numPr"):
                findings.append({"code": "template_block_numbering_failed", "artifact": key, "evidence": str(block.get("id") or "")})

        if not document.paragraphs:
            findings.append({"code": "template_direct_format_missing", "artifact": key, "evidence": "document has no paragraphs"})
            continue
        # A canonical draft produced by the public fixture may intentionally
        # omit the contact header.  When the header is present (the real
        # material contract), its two runs must be exact template copies.
        first = document.paragraphs[0]
        header_present = len(document.paragraphs) > 1 and ("|" in document.paragraphs[1].text or "@" in document.paragraphs[1].text)
        title_candidate = len(first.text.strip()) <= 80 and first.style.name == "Normal"
        if header_present or title_candidate:
            if rpr_signature(first.runs[0] if first.runs else None) != prototype_rpr_signature("title"):
                findings.append({"code": "template_title_format_failed", "artifact": key, "evidence": "title run differs from lane master"})
        if header_present and rpr_signature(document.paragraphs[1].runs[0] if document.paragraphs[1].runs else None) != prototype_rpr_signature("contact"):
            findings.append({"code": "template_contact_format_failed", "artifact": key, "evidence": "contact run differs from lane master"})

        if material == "cv":
            section = next((p for p in document.paragraphs if p.style.name == "Resume Section" and p.runs), None)
            if section is not None and rpr_signature(section.runs[0]) != prototype_rpr_signature("section"):
                findings.append({"code": "template_section_format_failed", "artifact": key, "evidence": "Resume Section direct formatting missing"})
            job = next((p for p in document.paragraphs if p.style.name == "Job Heading" and p.runs), None)
            if job is not None and rpr_signature(job.runs[0]) != prototype_rpr_signature("job_heading", 0):
                findings.append({"code": "template_job_heading_format_failed", "artifact": key, "evidence": "Job Heading direct formatting missing"})
            if prototype_has_ppr_tag("job_heading", "tabs") and job is not None and not has_ppr_tag(job, "tabs"):
                findings.append({"code": "template_job_heading_tabs_failed", "artifact": key, "evidence": "right-aligned date tab missing"})
            bullet = next((p for p in document.paragraphs if p.style.name == "Resume Bullet"), None)
            if prototype_has_ppr_tag("bullet", "numPr") and bullet is not None and not has_ppr_tag(bullet, "numPr"):
                findings.append({"code": "template_bullet_numbering_failed", "artifact": key, "evidence": "Resume Bullet numbering missing"})
            core = next((p for p in document.paragraphs if p.style.name == "Compact Line" and len(p.runs) >= 2), None)
            if core is not None and rpr_signature(core.runs[0]) != prototype_rpr_signature("core", 0):
                findings.append({"code": "template_core_format_failed", "artifact": key, "evidence": "Core Expertise label formatting missing"})
        else:
            subject = next((p for p in document.paragraphs if p.text.lstrip().casefold().startswith("re:") and p.runs), None)
            if subject is not None and rpr_signature(subject.runs[0]) != prototype_rpr_signature("subject"):
                findings.append({"code": "template_subject_format_failed", "artifact": key, "evidence": "subject direct formatting missing"})
            bullet = next((p for p in document.paragraphs if p.style.name == "Letter Bullet" and len(p.runs) >= 2), None)
            if bullet is not None and rpr_signature(bullet.runs[0]) != prototype_rpr_signature("bullet", 0):
                findings.append({"code": "template_letter_bullet_format_failed", "artifact": key, "evidence": "Cover Letter pillar label formatting missing"})
    return findings


def mechanical_format_gate(package: Path, workspace: Path) -> dict[str, Any]:
    """Final deterministic gate: format and derivation only, never semantics."""

    package = Path(package)
    names = expected_filenames(package, workspace)
    findings: list[dict[str, str]] = []
    receipt = _load(package / RENDER_RECEIPT_NAME)
    recorded_docx = receipt.get("docx_hashes") if isinstance(receipt.get("docx_hashes"), dict) else {}
    live_docx = {
        key: container_hash(package / names[key])
        for key in ("cv_docx", "cl_docx")
        if (package / names[key]).is_file()
    }
    if recorded_docx and recorded_docx != live_docx:
        findings.append({
            "code": "render_receipt_hash_mismatch",
            "artifact": "docx",
            "evidence": "render receipt does not match the live DOCX; rerender through the fixed gateway",
        })
    elif not recorded_docx and any(path.is_file() for path in (package / names["cv_docx"], package / names["cl_docx"])):
        findings.append({
            "code": "render_receipt_missing",
            "artifact": "docx",
            "evidence": "DOCX exists without a renderer receipt",
        })
    # The vNext artifact receipt is written only after DOCX/PDF format passes;
    # an older receipt must never be allowed to describe a new generation.
    for artifact_receipt_path in (
        package / "materials_vnext" / "artifact_hashes.json",
        package / "artifact_hashes.json",
    ):
        artifact_receipt = _load(artifact_receipt_path)
        recorded_files = artifact_receipt.get("files") if isinstance(artifact_receipt.get("files"), dict) else {}
        if not recorded_files:
            continue
        live_files = {
            name: container_hash(package / name)
            for name in recorded_files
            if (package / name).is_file()
        }
        if live_files != recorded_files:
            findings.append({
                "code": "artifact_receipt_hash_mismatch",
                "artifact": artifact_receipt_path.name,
                "evidence": "artifact hash receipt does not match the live outbound files; rerun render/PDF/format",
            })
    for key in ("cv_docx", "cl_docx", "cv_pdf", "cl_pdf"):
        path = package / names[key]
        if not path.is_file():
            findings.append({"code": "required_outbound_missing", "artifact": key, "evidence": names[key]})
            continue
        if path.name != names[key]:
            findings.append({"code": "filename_mismatch", "artifact": key, "evidence": path.name})
        for code in metadata_violations(path):
            findings.append({"code": code, "artifact": key, "evidence": path.name})
        if key.endswith("_pdf"):
            pages, text = _pdf_stats(path)
            if pages != 1:
                findings.append({"code": "page_count_exceeded", "artifact": key, "evidence": str(pages)})
            if len(text.strip()) < 80:
                findings.append({"code": "pdf_text_layer_missing", "artifact": key, "evidence": str(len(text.strip()))})
    for code in _render_matches_canonical(package, names):
        findings.append({"code": code, "artifact": "docx", "evidence": "canonical text not preserved"})
    findings.extend(_template_style_findings(package, names))
    report = {
        "schema_version": 1,
        "scope": "mechanical_only",
        "checks": [
            "page_count",
            "text_layer",
            "filename",
            "metadata",
            "canonical_render_equivalence",
            "template_style_contract",
            "template_direct_format_contract",
        ],
        "status": "passed" if not findings else "failed",
        "format_passed": not findings,
        "findings": findings,
        "filenames": names,
        "layout_balance": _load(Path(package) / RENDER_RECEIPT_NAME).get("layout_balance") or {},
        "checked_at": _now(),
    }
    atomic_write_json(package / FORMAT_REPORT_NAME, report)
    return report
