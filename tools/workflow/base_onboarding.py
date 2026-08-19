"""Product-owned onboarding for lane CV/Cover Letter masters.

This module is deliberately runtime-facing: it writes only below the selected
private workspace (normally ``JobSearch_2026``).  The product repository ships
the schema and the style contract, never a candidate's filled CV/CL.

The important boundary is:

    user facts -> model structured base response -> deterministic checks
    -> anonymous product renderer -> preview -> explicit activation

An unreviewed draft is named ``draft_*`` and therefore cannot be selected by
the application-material renderer, which only accepts activated ``master_*`` /
``cl_master_*`` files.
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from tools.job_materials.paths import (
    find_latest_cl_master_docx,
    find_latest_master_docx,
    lane_masters_folder,
    load_lanes,
    profile_dir,
)


BASE_SCHEMA_VERSION = 1
BASE_ENGINE_VERSION = "base-onboarding-v1"
BASE_FORMAT_CONTRACT_VERSION = "jobsflow-base-format-v1"
BASE_REQUEST_DIR = "base_requests"
BASE_STATE_DIR = "bases_runtime"

# This is a product contract, not a personal template.  The renderer remains
# the authority for per-job output; these values make first-run masters use
# the same stable, ATS-friendly visual family.
FORMAT_CONTRACT: dict[str, Any] = {
    "version": BASE_FORMAT_CONTRACT_VERSION,
    "page": {"size": "A4", "single_column": True, "text_layer_required": True},
    "cv": {
        "styles": ["Resume Section", "Job Heading", "Resume Bullet", "Compact Line"],
        "body_font": "Calibri",
        "body_size_pt": 10.5,
        "heading_color": "17365D",
        "job_heading_color": "1C1C1C",
        "heading_bold": True,
        "max_pages": 1,
        "required_sections": ["Summary", "Core Expertise", "Experience", "Education"],
    },
    "cover_letter": {
        "styles": ["Letter Body", "Letter Bullet", "Letter Compact"],
        "body_font": "Calibri",
        "body_size_pt": 10.5,
        "heading_color": "17365D",
        "heading_bold": True,
        "max_pages": 1,
        "required_parts": ["opening", "evidence", "closing"],
    },
    "prohibitions": [
        "Do not use a blank document or a model-selected renderer.",
        "Do not put personal facts into the tracked product repository.",
        "Do not activate a base before deterministic validation and user confirmation.",
    ],
}


def _now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _json(path: Path, default: Any = None) -> Any:
    try:
        return json.loads(Path(path).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError, TypeError):
        return default


def _write(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _sha(path: Path) -> str:
    h = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _candidate_name(root: Path) -> str:
    config = _json(profile_dir(root) / "config.personal.json", {})
    return str((config or {}).get("candidate_name") or (config or {}).get("name") or "Candidate").strip()


def _contact_line(root: Path) -> str:
    config = _json(profile_dir(root) / "config.personal.json", {})
    if not isinstance(config, dict):
        return "Contact details"
    values = [str(config.get(key) or "").strip() for key in ("email", "phone", "location")]
    return " | ".join(value for value in values if value) or "Contact details"


def _facts(root: Path) -> list[dict[str, Any]]:
    payload = _json(profile_dir(root) / "fact_evidence.json", {})
    records = payload.get("records") if isinstance(payload, dict) else []
    if not isinstance(records, list):
        return []
    return [item for item in records if isinstance(item, dict) and item.get("evidence_id") and item.get("claim")]


def _lane_keys(root: Path) -> list[str]:
    return sorted(str(key).upper() for key in load_lanes(root).keys())


def request_path(root: Path, lane: str) -> Path:
    return profile_dir(root) / BASE_REQUEST_DIR / str(lane).upper() / "request.json"


def response_path(root: Path, lane: str) -> Path:
    return profile_dir(root) / BASE_REQUEST_DIR / str(lane).upper() / "response.json"


def state_path(root: Path, lane: str) -> Path:
    return profile_dir(root) / BASE_STATE_DIR / f"{str(lane).upper()}_onboarding.json"


def _fact_payload(root: Path) -> list[dict[str, str]]:
    return [
        {
            "evidence_id": str(item.get("evidence_id")),
            "claim": " ".join(str(item.get("claim") or "").split()),
            "status": str(item.get("status") or "user_imported"),
        }
        for item in _facts(root)
    ]


def write_request(root: Path, lane: str) -> Path:
    lane = str(lane).upper()
    lanes = load_lanes(root)
    meta = lanes.get(lane)
    if not meta:
        raise ValueError(f"unknown_lane:{lane}")
    facts = _fact_payload(root)
    request = {
        "schema_version": BASE_SCHEMA_VERSION,
        "artifact_type": "jobsflow_base_request",
        "engine_version": BASE_ENGINE_VERSION,
        "lane": lane,
        "lane_label": str(meta.get("label") or lane),
        "lane_emphasis": [item for item in str(meta.get("emphasis") or "").split(",") if item],
        "candidate_name": _candidate_name(root),
        "source_facts": facts,
        "source_fact_count": len(facts),
        "response_path": str(response_path(root, lane).resolve()),
        "required_output": {
            "artifact_type": "jobsflow_base_response",
            "lane": lane,
            "cv": {
                "summary": {"text": "...", "evidence_ids": ["EVID-..."]},
                "core": [{"text": "...", "evidence_ids": ["EVID-..."]}],
                "experience": [{"heading": "Employer — Role | Dates", "bullets": [{"text": "Action + object + method/result", "evidence_ids": ["EVID-..."]}]}],
                "education": [{"text": "Degree / institution / date", "evidence_ids": ["EVID-..."]}],
                "qualifications": [{"text": "Language / certification / qualification", "evidence_ids": ["EVID-..."]}],
            },
            "cover_letter": {
                "opening": "A concise lane-specific opening grounded in the profile.",
                "pillars": ["Evidence-backed value point 1", "Evidence-backed value point 2"],
                "closing": "A concise general closing; no company is hard-coded.",
            },
        },
        "model_contract": [
            "Use only source_facts or explicitly user-confirmed profile information.",
            "Do not invent employers, duties, tools, credentials, metrics, dates or outcomes.",
            "CV and Cover Letter are parallel masters; do not derive one from the other.",
            "Keep the response structured and concise. The host owns DOCX/PDF formatting.",
        ],
        "format_contract": FORMAT_CONTRACT,
        "created_at": _now(),
    }
    path = request_path(root, lane)
    _write(path, request)
    return path


def prepare_requests(root: Path, lanes: list[str] | None = None) -> list[Path]:
    selected = [str(item).upper() for item in (lanes or _lane_keys(root))]
    return [write_request(root, lane) for lane in selected]


def _normalise_list(value: Any) -> list[Any]:
    if value is None:
        return []
    return value if isinstance(value, list) else [value]


def _text(value: Any) -> str:
    return " ".join(str(value or "").split()).strip()


def _evidence_ids(value: Any) -> list[str]:
    if isinstance(value, dict):
        value = value.get("evidence_ids") or value.get("source_evidence_ids") or []
    return [str(item).strip() for item in _normalise_list(value) if str(item).strip()]


def _block(value: Any) -> tuple[str, list[str]]:
    if isinstance(value, dict):
        return _text(value.get("text") or value.get("content")), _evidence_ids(value)
    return _text(value), []


def _normalise_response(payload: dict[str, Any], lane: str) -> dict[str, Any]:
    if str(payload.get("artifact_type") or "") != "jobsflow_base_response":
        raise ValueError("base_response_artifact_type_invalid")
    if str(payload.get("lane") or "").upper() != lane:
        raise ValueError("base_response_lane_mismatch")
    cv = payload.get("cv") if isinstance(payload.get("cv"), dict) else {}
    cl = payload.get("cover_letter") if isinstance(payload.get("cover_letter"), dict) else {}
    summary = _block(cv.get("summary"))
    core = [_block(item) for item in _normalise_list(cv.get("core"))]
    experience: list[dict[str, Any]] = []
    for item in _normalise_list(cv.get("experience")):
        if not isinstance(item, dict):
            continue
        heading = _text(item.get("heading") or item.get("title"))
        bullets = [_block(value) for value in _normalise_list(item.get("bullets"))]
        experience.append({"heading": heading, "bullets": bullets, "evidence_ids": _evidence_ids(item)})
    education = [_block(item) for item in _normalise_list(cv.get("education"))]
    qualifications = [_block(item) for item in _normalise_list(cv.get("qualifications"))]
    opening = _block(cl.get("opening"))
    pillars = [_block(item) for item in _normalise_list(cl.get("pillars") or cl.get("evidence"))]
    closing = _block(cl.get("closing"))
    return {
        "schema_version": BASE_SCHEMA_VERSION,
        "artifact_type": "jobsflow_base_response",
        "lane": lane,
        "candidate_name": _text(payload.get("candidate_name")),
        "cv": {
            "summary": summary,
            "core": core,
            "experience": experience,
            "education": education,
            "qualifications": qualifications,
        },
        "cover_letter": {"opening": opening, "pillars": pillars, "closing": closing},
    }


def _validate_content(response: dict[str, Any], root: Path) -> list[str]:
    errors: list[str] = []
    facts = _facts(root)
    fact_ids = {str(item.get("evidence_id")) for item in facts}
    fact_blob = " ".join(str(item.get("claim") or "") for item in facts).casefold()
    fact_numbers = set(re.findall(r"\d+(?:[.,]\d+)?", fact_blob))
    forbidden = ("[company", "[role", "[date", "insert ", "lorem ipsum", "rather than", "lack of", "no experience")
    texts: list[tuple[str, list[str], str]] = []

    def add(label: str, pair: tuple[str, list[str]], kind: str = "content") -> None:
        value, ids = pair
        if not value:
            errors.append(f"{label}_missing")
            return
        if any(token in value.casefold() for token in forbidden):
            errors.append(f"{label}_placeholder_or_negative")
        claim_numbers = set(re.findall(r"\d+(?:[.,]\d+)?", value))
        if claim_numbers - fact_numbers:
            errors.append(f"{label}_numeric_claim_without_facts")
        unknown = [item for item in ids if item not in fact_ids]
        if unknown:
            errors.append(f"{label}_unknown_evidence:{','.join(unknown)}")
        texts.append((value, ids, kind))

    cv = response["cv"]
    add("cv_summary", cv["summary"])
    for index, item in enumerate(cv["core"]):
        add(f"cv_core_{index + 1}", item)
    for index, item in enumerate(cv["experience"]):
        if not item["heading"]:
            errors.append(f"experience_{index + 1}_heading_missing")
        if not item["bullets"]:
            errors.append(f"experience_{index + 1}_bullets_missing")
        for bindex, bullet in enumerate(item["bullets"]):
            add(f"experience_{index + 1}_bullet_{bindex + 1}", bullet, "bullet")
            # STAR's hard minimum: a bullet must contain an action and object;
            # method/result is encouraged but cannot be invented by the gate.
            words = set(re.findall(r"[a-zA-Z]{3,}", bullet[0].casefold()))
            if len(words) < 4:
                errors.append(f"experience_{index + 1}_bullet_{bindex + 1}_too_thin")
    for index, item in enumerate(cv["education"]):
        add(f"education_{index + 1}", item)
    for index, item in enumerate(cv["qualifications"]):
        add(f"qualification_{index + 1}", item)
    if not cv["core"]:
        errors.append("cv_core_missing")
    if not cv["experience"] and not cv["education"]:
        errors.append("cv_experience_or_education_missing")

    cl = response["cover_letter"]
    add("cl_opening", cl["opening"])
    for index, item in enumerate(cl["pillars"]):
        add(f"cl_pillar_{index + 1}", item)
    add("cl_closing", cl["closing"])
    if not cl["pillars"]:
        errors.append("cl_evidence_missing")
    if len(cl["pillars"]) > 3:
        errors.append("cl_pillars_over_three")

    if not facts:
        errors.append("profile_facts_missing")
    # Unreferenced blocks are allowed for generic headings/closing language,
    # but substantive bullets must either point at a fact or overlap the fact
    # store enough to be reviewable by a low-capability model.
    for label, value, ids, kind in (
        ("content", value, ids, kind) for value, ids, kind in texts
    ):
        if kind == "bullet" and not ids:
            overlap = sum(1 for token in re.findall(r"[a-zA-Z]{4,}", value.casefold()) if token in fact_blob)
            if overlap < 2:
                errors.append("bullet_without_evidence_anchor")
    return sorted(set(errors))


def validate_response(root: Path, lane: str, content_path: Path) -> tuple[dict[str, Any] | None, list[str]]:
    raw = _json(Path(content_path))
    if not isinstance(raw, dict):
        return None, ["base_response_invalid_json"]
    try:
        response = _normalise_response(raw, str(lane).upper())
    except ValueError as exc:
        return None, [str(exc)]
    errors = _validate_content(response, root)
    return response, errors


def _set_style(style, *, font: str = "Calibri", size: float = 10.5, color: str | None = None, bold: bool = False) -> None:
    from docx.shared import Pt, RGBColor

    style.font.name = font
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def _new_document(material: str):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.27 if material == "cv" else 1.91)
        section.bottom_margin = Cm(1.27 if material == "cv" else 1.91)
        section.left_margin = Cm(1.83 if material == "cv" else 2.29)
        section.right_margin = Cm(1.83 if material == "cv" else 2.29)
    normal = doc.styles["Normal"]
    _set_style(normal)
    styles = {
        "Resume Section": (10.5, "17365D", True),
        "Job Heading": (10.5, "1C1C1C", True),
        "Resume Bullet": (10.5, None, False),
        "Compact Line": (10.5, None, False),
        "Letter Body": (10.5, None, False),
        "Letter Bullet": (10.5, None, False),
        "Letter Compact": (10.5, None, False),
    }
    for name, (size, color, bold) in styles.items():
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, 1)  # WD_STYLE_TYPE.PARAGRAPH
        _set_style(style, size=size, color=color, bold=bold)
    return doc


def _p(doc, value: str, style: str = "Normal", *, bold: bool = False, color: str | None = None) -> None:
    from docx.shared import RGBColor

    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(value)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _render_base_docs(root: Path, lane: str, response: dict[str, Any]) -> dict[str, Path]:
    folder = lane_masters_folder(lane, root)
    if folder is None:
        raise ValueError(f"lane_folder_missing:{lane}")
    folder.mkdir(parents=True, exist_ok=True)
    candidate = response.get("candidate_name") or _candidate_name(root)
    cv_doc = _new_document("cv")
    _p(cv_doc, candidate, "Normal", bold=True, color="17365D")
    _p(cv_doc, _contact_line(root), "Normal")
    _p(cv_doc, "Summary", "Resume Section")
    _p(cv_doc, response["cv"]["summary"][0], "Normal")
    _p(cv_doc, "Core Expertise", "Resume Section")
    for item, _ids in response["cv"]["core"]:
        _p(cv_doc, item, "Compact Line")
    _p(cv_doc, "Experience", "Resume Section")
    for item in response["cv"]["experience"]:
        _p(cv_doc, item["heading"], "Job Heading")
        for bullet, _ids in item["bullets"]:
            _p(cv_doc, bullet, "Resume Bullet")
    if response["cv"]["education"]:
        _p(cv_doc, "Education", "Resume Section")
        for item, _ids in response["cv"]["education"]:
            _p(cv_doc, item, "Compact Line")
    if response["cv"]["qualifications"]:
        _p(cv_doc, "Qualifications", "Resume Section")
        for item, _ids in response["cv"]["qualifications"]:
            _p(cv_doc, item, "Compact Line")

    cl_doc = _new_document("cover_letter")
    _p(cl_doc, candidate, "Normal", bold=True, color="17365D")
    _p(cl_doc, _contact_line(root), "Normal")
    _p(cl_doc, "Re: Application", "Normal", bold=True, color="17365D")
    _p(cl_doc, "Hiring Team", "Letter Compact")
    _p(cl_doc, response["cover_letter"]["opening"][0], "Letter Body")
    for item, _ids in response["cover_letter"]["pillars"]:
        _p(cl_doc, item, "Letter Bullet")
    _p(cl_doc, response["cover_letter"]["closing"][0], "Letter Body")
    _p(cl_doc, "Yours faithfully,", "Letter Compact")
    _p(cl_doc, candidate, "Letter Compact")

    cv_path = folder / f"draft_master_{lane}_v0.1.docx"
    cl_path = folder / f"draft_cl_master_{lane}_v0.1.docx"
    cv_doc.save(str(cv_path))
    cl_doc.save(str(cl_path))
    return {"cv": cv_path, "cover_letter": cl_path}


def _docx_errors(path: Path, material: str) -> list[str]:
    try:
        from docx import Document

        doc = Document(str(path))
    except Exception as exc:
        return [f"docx_unreadable:{material}:{exc}"]
    required = FORMAT_CONTRACT[material]["styles"]
    available = {str(style.name) for style in doc.styles}
    errors = [f"style_missing:{material}:{name}" for name in required if name not in available]
    if not any(str(p.text).strip() for p in doc.paragraphs):
        errors.append(f"docx_empty:{material}")
    expected_margin = (1.27, 1.27, 1.83, 1.83) if material == "cv" else (1.91, 1.91, 2.29, 2.29)
    if doc.sections:
        section = doc.sections[0]
        actual_margin = tuple(round(float(value.cm), 2) for value in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin))
        if actual_margin != expected_margin:
            errors.append(f"page_margin_contract_failed:{material}")
    for name in ("Resume Section", "Job Heading") if material == "cv" else ():
        if name not in available:
            continue
        style = doc.styles[name]
        sample = next((paragraph for paragraph in doc.paragraphs if str(paragraph.style.name) == name), None)
        sample_run = next((run for run in (sample.runs if sample is not None else []) if str(run.text or "").strip()), None)
        color_value = style.font.color.rgb
        if not color_value and sample_run is not None:
            color_value = sample_run.font.color.rgb
        color = str(color_value or "").replace("#", "").upper()
        allowed_colors = {"17365D", "17365D00"} if name == "Resume Section" else {"1C1C1C", "1C1C1C00"}
        if color and color not in allowed_colors:
            errors.append(f"heading_color_contract_failed:{material}:{name}")
        bold = style.font.bold
        if bold is None and sample_run is not None:
            bold = sample_run.bold
        if bold is not True:
            errors.append(f"heading_weight_contract_failed:{material}:{name}")
    return errors


def status(root: Path, lane: str | None = None) -> dict[str, Any]:
    lanes = [str(lane).upper()] if lane else _lane_keys(root)
    results = []
    for item in lanes:
        folder = lane_masters_folder(item, root)
        cv = find_latest_master_docx(item, root)
        cl = find_latest_cl_master_docx(item, root)
        state = _json(state_path(root, item), {})
        request = request_path(root, item)
        ready = bool(cv and cl)
        blockers = []
        if not request.exists() and not ready:
            blockers.append("base_request_missing")
        if not cv:
            blockers.append("baseline_cv_master_missing")
        if not cl:
            blockers.append("baseline_cl_master_missing")
        if cv:
            blockers.extend(_docx_errors(cv, "cv"))
        if cl:
            blockers.extend(_docx_errors(cl, "cover_letter"))
        results.append(
            {
                "lane": item,
                "folder": str(folder) if folder else "",
                "ready": ready and not blockers,
                "status": str(state.get("status") or ("active_existing_master" if ready else "pending")),
                "cv_master": str(cv) if cv else "",
                "cl_master": str(cl) if cl else "",
                "blockers": sorted(set(blockers)),
            }
        )
    return {
        "schema_version": BASE_SCHEMA_VERSION,
        "engine_version": BASE_ENGINE_VERSION,
        "format_contract": BASE_FORMAT_CONTRACT_VERSION,
        "lanes": results,
        "ready": bool(results) and all(bool(item["ready"]) for item in results),
        "next_action": "continue" if results and all(bool(item["ready"]) for item in results) else "base_init_or_generate_and_confirm",
    }


def init(root: Path, lane: str | None = None) -> dict[str, Any]:
    paths = prepare_requests(root, [str(lane).upper()] if lane else None)
    return {"status": "initialized", "requests": [str(path) for path in paths], "next_action": "fill_response_json"}


def generate(root: Path, lane: str, content: Path) -> dict[str, Any]:
    lane = str(lane).upper()
    expected = response_path(root, lane).resolve()
    supplied = Path(content).expanduser().resolve()
    if supplied != expected:
        return {"status": "blocked", "lane": lane, "blockers": ["base_response_path_invalid"], "expected_response": str(expected)}
    response, errors = validate_response(root, lane, supplied)
    if errors or response is None:
        return {"status": "blocked", "lane": lane, "blockers": errors or ["base_response_invalid"]}
    paths = _render_base_docs(root, lane, response)
    state = {
        "schema_version": BASE_SCHEMA_VERSION,
        "engine_version": BASE_ENGINE_VERSION,
        "lane": lane,
        "status": "draft_pending_review",
        "response_sha256": _sha(supplied),
        "drafts": {key: str(path) for key, path in paths.items()},
        "draft_sha256": {key: _sha(path) for key, path in paths.items()},
        "format_contract": BASE_FORMAT_CONTRACT_VERSION,
        "created_at": _now(),
    }
    _write(state_path(root, lane), state)
    return {"status": "drafted", "lane": lane, "drafts": state["drafts"], "next_action": "base confirm --lane " + lane}


def confirm(root: Path, lane: str, *, confirmed: bool = False) -> dict[str, Any]:
    lane = str(lane).upper()
    state = _json(state_path(root, lane), {})
    drafts = state.get("drafts") if isinstance(state, dict) else {}
    cv = Path(str((drafts or {}).get("cv") or ""))
    cl = Path(str((drafts or {}).get("cover_letter") or ""))
    errors = []
    supplied = response_path(root, lane)
    if not supplied.is_file():
        errors.append("base_response_missing")
    elif state.get("response_sha256") and state.get("response_sha256") != _sha(supplied):
        errors.append("base_response_changed_after_generate")
    stored_hashes = state.get("draft_sha256") if isinstance(state.get("draft_sha256"), dict) else {}
    for key, path in (("cv", cv), ("cover_letter", cl)):
        if path.is_file() and stored_hashes.get(key) and stored_hashes.get(key) != _sha(path):
            errors.append(f"base_draft_changed:{key}")
    if not errors and supplied.is_file():
        _response, response_errors = validate_response(root, lane, supplied)
        errors.extend(response_errors)
    errors.extend(_docx_errors(cv, "cv") if cv.is_file() else ["base_cv_draft_missing"])
    errors.extend(_docx_errors(cl, "cover_letter") if cl.is_file() else ["base_cl_draft_missing"])
    if errors:
        return {"status": "blocked", "lane": lane, "blockers": sorted(set(errors))}
    if not confirmed:
        return {"status": "preview", "lane": lane, "drafts": drafts, "next_action": "repeat_with_--confirm"}
    folder = lane_masters_folder(lane, root)
    if folder is None:
        return {"status": "blocked", "lane": lane, "blockers": ["lane_folder_missing"]}
    folder.mkdir(parents=True, exist_ok=True)
    active_cv = folder / f"master_{lane}_v1.0.docx"
    active_cl = folder / f"cl_master_{lane}_v1.0.docx"
    for source, target in ((cv, active_cv), (cl, active_cl)):
        if target.exists():
            history = folder / ".history" / f"base-{datetime.now().strftime('%Y%m%d%H%M%S')}"
            history.mkdir(parents=True, exist_ok=True)
            shutil.move(str(target), str(history / target.name))
        shutil.move(str(source), str(target))
    state.update({"status": "active", "active_cv": str(active_cv), "active_cl": str(active_cl), "activated_at": _now(), "active_sha256": {"cv": _sha(active_cv), "cover_letter": _sha(active_cl)}})
    _write(state_path(root, lane), state)
    return {"status": "activated", "lane": lane, "cv_master": str(active_cv), "cl_master": str(active_cl), "format_contract": BASE_FORMAT_CONTRACT_VERSION}


def handle(root: Path, command: str, lane: str = "", content: Path | None = None, confirmed: bool = False) -> dict[str, Any]:
    command = str(command or "status")
    if command == "status":
        return status(root, lane or None)
    if command == "init":
        return init(root, lane or None)
    if command == "generate":
        if not lane or content is None:
            return {"status": "blocked", "blockers": ["base_lane_and_response_required"]}
        return generate(root, lane, content)
    if command == "confirm":
        if not lane:
            return {"status": "blocked", "blockers": ["base_lane_required"]}
        return confirm(root, lane, confirmed=confirmed)
    return {"status": "blocked", "blockers": [f"base_command_unknown:{command}"]}
